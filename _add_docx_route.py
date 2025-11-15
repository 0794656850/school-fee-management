from pathlib import Path
app_path = Path('app.py')
s = app_path.read_text(encoding='utf-8')
if 'def export_school_profile_docx()' not in s:
    block = '''

# ---------- SCHOOL PROFILE DOCX ----------
@app.route("/export_school_profile_docx")
def export_school_profile_docx():
    # Pro gating consistent with other exports
    if not is_pro_enabled(app):
        try:
            flash('Exports are available in Pro. Please upgrade.', 'warning')
        except Exception:
            pass
        return redirect(url_for('monetization.index'))

    try:
        from docx import Document
        from docx.shared import Pt, Inches
    except Exception:
        try:
            flash('Word export requires python-docx. Please install python-docx.', 'warning')
        except Exception:
            pass
        return redirect(url_for('reports'))

    db = get_db_connection()
    cur = db.cursor(dictionary=True)
    sid = session.get('school_id')

    # School profile (from settings)
    school_name = get_setting('SCHOOL_NAME') or get_setting('APP_NAME') or 'School'
    school_address = get_setting('SCHOOL_ADDRESS') or ''
    school_phone = get_setting('SCHOOL_PHONE') or ''
    school_email = get_setting('SCHOOL_EMAIL') or ''

    # Totals
    cur.execute("SELECT COUNT(*) AS total FROM students WHERE school_id=%s", (sid,))
    total_students = (cur.fetchone() or {}).get('total', 0)
    cur.execute("SELECT COALESCE(SUM(amount),0) AS total FROM payments WHERE method <> 'Credit Transfer' AND school_id=%s", (sid,))
    total_collected = (cur.fetchone() or {}).get('total', 0)
    cur.execute("SHOW COLUMNS FROM students LIKE 'balance'")
    has_balance = bool(cur.fetchone())
    column = 'balance' if has_balance else 'fee_balance'
    cur.execute(f"SELECT COALESCE(SUM({column}),0) AS total FROM students WHERE school_id=%s", (sid,))
    total_balance = (cur.fetchone() or {}).get('total', 0)
    cur.execute("SELECT COALESCE(SUM(credit),0) AS total FROM students WHERE school_id=%s", (sid,))
    total_credit = (cur.fetchone() or {}).get('total', 0)

    # Per-class summary
    cur.execute("""
        SELECT class_name AS class,
               COUNT(*) AS total_students,
               COALESCE(SUM(COALESCE(balance, fee_balance)),0) AS total_pending,
               COALESCE(SUM(credit),0) AS total_credit
        FROM students WHERE school_id=%s GROUP BY class_name ORDER BY class_name
    """, (sid,))
    class_rows = cur.fetchall() or []
    cur.execute("""
        SELECT s.class_name AS class, COALESCE(SUM(p.amount),0) AS total_paid
        FROM payments p JOIN students s ON s.id=p.student_id
        WHERE p.school_id=%s AND p.method <> 'Credit Transfer' GROUP BY s.class_name
    """, (sid,))
    paid_map = {r['class']: float(r.get('total_paid') or 0) for r in (cur.fetchall() or [])}

    db.close()

    # Build document
    doc = Document()
    doc.add_heading(f"{school_name} — Fee Report", 0)
    p = doc.add_paragraph()
    p.add_run(f"Address: {school_address}\n").font.size = Pt(10)
    p.add_run(f"Phone: {school_phone}  Email: {school_email}\n").font.size = Pt(10)

    from datetime import datetime as _dt
    doc.add_paragraph(f"Generated: {_dt.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # Summary
    doc.add_heading('Summary', level=1)
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light List'
    labels = ['Total Students','Total Collected (KES)','Total Pending (KES)','Total Credit (KES)','Collection Rate (%)']
    try:
        rate = (float(total_collected) / (float(total_collected) + float(total_balance)) * 100.0) if (float(total_collected)+float(total_balance))>0 else 0.0
    except Exception:
        rate = 0.0
    values = [total_students, total_collected, total_balance, total_credit, round(rate,1)]
    for i,(k,v) in enumerate(zip(labels, values)):
        table.cell(i,0).text = str(k)
        table.cell(i,1).text = f"{v}"

    # Class Summary
    doc.add_heading('Class Summary', level=1)
    ct = doc.add_table(rows=1, cols=6)
    ct.style = 'Light Grid'
    hdr = ['Class','Total Students','Paid (KES)','Pending (KES)','Credit (KES)','Rate %']
    for j,h in enumerate(hdr):
        ct.rows[0].cells[j].text = h
    for row in class_rows:
        paid = float(paid_map.get(row.get('class') or '', 0))
        pending = float(row.get('total_pending') or 0)
        credit = float(row.get('total_credit') or 0)
        total = paid + pending
        rate = round((paid/total*100.0),1) if total>0 else 0.0
        r = ct.add_row().cells
        r[0].text = str(row.get('class') or '')
        r[1].text = str(row.get('total_students') or 0)
        r[2].text = str(round(paid,2))
        r[3].text = str(round(pending,2))
        r[4].text = str(round(credit,2))
        r[5].text = str(rate)

    # Return .docx
    from io import BytesIO
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    ts = _dt.now().strftime('%Y%m%d_%H%M%S')
    return Response(bio.getvalue(), mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', headers={'Content-Disposition': f'attachment; filename=school_profile_{ts}.docx'})
'''
    s += block
    app_path.write_text(s, encoding='utf-8')
    print('Added export_school_profile_docx route')
else:
    print('export_school_profile_docx already present')
