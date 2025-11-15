from flask import Flask, render_template, request, redirect, url_for, jsonify, flash, Response, session
import json
import hmac
import hashlib
import mysql.connector
from datetime import datetime
import csv
from io import StringIO
from urllib.parse import urlparse
import os
from config import Config
from routes.reminder_routes import reminder_bp
from routes.admin_routes import admin_bp
from routes.auth_routes import auth_bp
from routes.credit_routes import credit_bp
from routes.credit_routes import ensure_credit_ops_table, ensure_students_credit_column, ensure_credit_transfers_table
from routes.term_routes import (
    term_bp,
    ensure_academic_terms_table,
    ensure_payments_term_columns,
    get_or_seed_current_term,
    ensure_student_enrollments_table,
    ensure_term_fees_table,
)
# Extend term routes (bulk flat fees)
import routes.term_flat_routes  # noqa: F401 - registers extra routes on term_bp
from utils.notify import normalize_phone
# Gmail API is optional; import lazily/fail-safe to avoid hard import dependency
try:
    from utils.gmail_api import (
        send_email as gmail_send_email,
        send_email_html as gmail_send_email_html,
    )
except Exception:  # pragma: no cover - optional dependency not installed
    def gmail_send_email(*args, **kwargs):  # type: ignore
        return False

    def gmail_send_email_html(*args, **kwargs):  # type: ignore
        return False
from routes.mpesa_routes import mpesa_bp
from routes.gmail_oauth_routes import gmail_oauth_bp
from routes.student_auth import student_auth_bp
from routes.student_portal import student_portal_bp
from routes.student_auth import ensure_student_portal_columns
from utils.security import hash_password
from routes.guardian_routes import guardian_bp
from routes.newsletter_routes import newsletter_bp, ensure_newsletters_table
from extensions import db, migrate
from billing import billing_bp
from routes.defaulter_routes import recovery_bp
from utils.settings import get_setting, set_school_setting
from utils.users import ensure_user_tables
from routes.ai_routes import ai_bp
from utils.audit import log_event
from utils.ledger import ensure_ledger_table, add_entry
from utils.tenant import (
    ensure_school_id_columns,
    ensure_schools_table,
    get_or_create_school,
    slugify_code,
    bootstrap_new_school,
    ensure_unique_indices_per_school,
)
from werkzeug.utils import secure_filename
try:
    from werkzeug.middleware.proxy_fix import ProxyFix
except Exception:  # pragma: no cover
    ProxyFix = None  # type: ignore

app = Flask(__name__)

# Load configuration from Config (falls back to sensible defaults inside Config)
app.config.from_object(Config)

# Trust reverse proxy headers for scheme/host when enabled
try:
    if (app.config.get("TRUST_PROXY", True)) and ProxyFix is not None:
        app.wsgi_app = ProxyFix(app.wsgi_app, x_proto=1, x_host=1)  # type: ignore[attr-defined]
except Exception:
    pass

# Initialize SQLAlchemy (for billing models) if available
try:
    from extensions import db as _db  # type: ignore
    try:
        _db.init_app(app)
        # Create billing tables if missing (limited to models imported below)
        try:
            with app.app_context():
                from billing import LicenseRequest, LicenseKey  # noqa: F401
                _db.create_all()
        except Exception:
            pass
    except Exception:
        pass
except Exception:
    pass

# Set modern security headers on every response (best effort, non-breaking)
@app.after_request
def _set_security_headers(resp):
    try:
        # Basic hardening
        resp.headers.setdefault("X-Content-Type-Options", "nosniff")
        resp.headers.setdefault("X-Frame-Options", "SAMEORIGIN")
        resp.headers.setdefault("Referrer-Policy", "strict-origin-when-cross-origin")
        resp.headers.setdefault("Permissions-Policy", "geolocation=(), microphone=(), camera=()")
        # CSP kept permissive because of inline styles/scripts and CDN usage
        # Tighten this once inline CSS/JS is externalized.
        csp = (
            "default-src 'self'; "
            "img-src 'self' data: blob: https:; "
            "style-src 'self' 'unsafe-inline' https://cdn.tailwindcss.com https://fonts.googleapis.com; "
            "font-src 'self' https://fonts.gstatic.com; "
            "script-src 'self' 'unsafe-inline' https://cdn.tailwindcss.com https://unpkg.com https://cdn.jsdelivr.net https://cdnjs.cloudflare.com; "
            "connect-src 'self' data: blob:; "
            "frame-ancestors 'self'"
        )
        resp.headers.setdefault("Content-Security-Policy", csp)
        # HSTS only when cookies marked secure (implies HTTPS)
        if app.config.get("SESSION_COOKIE_SECURE", False):
            resp.headers.setdefault("Strict-Transport-Security", "max-age=31536000; includeSubDomains")
    except Exception:
        # Do not block response if header setting fails
        pass
    return resp

# Audit trail removed: no per-request logging hook

# Allow HTTP OAuth redirect for local Gmail auth
# oauthlib enforces HTTPS by default; for localhost development we
# enable insecure transport only when an HTTP redirect is intended.
try:
    _gmail_uri = os.environ.get("GMAIL_REDIRECT_URI", "").strip()
    _prefer_http = (
        _gmail_uri.startswith("http://")
        or ("127.0.0.1" in _gmail_uri)
        or ("localhost" in _gmail_uri)
        or app.config.get("PREFERRED_URL_SCHEME", "http") == "http"
    )
    if _prefer_http and not os.environ.get("OAUTHLIB_INSECURE_TRANSPORT"):
        os.environ["OAUTHLIB_INSECURE_TRANSPORT"] = "1"
except Exception:
    # Nonâ€‘fatal; if this fails, user can set env manually
    pass

# Enforce HTTPS for all requests (except localhost) when enabled
@app.before_request
def _enforce_https_redirect():
    try:
        if not app.config.get("ENFORCE_HTTPS", False):
            return None
        # Skip for local development hosts
        host = (request.host or "").split(":")[0]
        if host in ("127.0.0.1", "localhost"):
            return None
        # Honor proxy headers if present
        xf_proto = (request.headers.get("X-Forwarded-Proto", "").split(",")[0].strip().lower())
        if request.is_secure or xf_proto == "https":
            return None
        # Redirect to HTTPS
        url = request.url.replace("http://", "https://", 1)
        return redirect(url, code=301)
    except Exception:
        return None

# Initialize optional extensions
try:
    from extensions import limiter, mail  # type: ignore
    try:
        limiter.init_app(app)  # no-op if fallback limiter
    except Exception:
        pass
    try:
        # Initialize Flask-Mail for SMTP fallback (safe even if unconfigured)
        mail.init_app(app)
    except Exception:
        pass
except Exception:
    # If extensions import fails, continue startup
    pass

# Ensure secret key is set (Config provides default). Keeping compatibility if env overrides.
app.secret_key = app.config.get("SECRET_KEY", os.environ.get("SECRET_KEY", "secret123"))

# Assign a per-request correlation id for tracing
try:
    import uuid
    from flask import g

    @app.before_request
    def _assign_request_id():
        try:
            g.request_id = uuid.uuid4().hex[:16]
        except Exception:
            pass
except Exception:
    pass

# Initialize database + migrations (for SQLAlchemy models like billing)
try:
    db.init_app(app)
    migrate.init_app(app, db)
except Exception:
    # Continue even if SQLAlchemy isn't fully configured
    pass

# Register blueprints
app.register_blueprint(reminder_bp)
app.register_blueprint(admin_bp)
app.register_blueprint(auth_bp)
app.register_blueprint(credit_bp)
app.register_blueprint(term_bp)
app.register_blueprint(mpesa_bp)
app.register_blueprint(student_auth_bp)
app.register_blueprint(student_portal_bp)
app.register_blueprint(guardian_bp)
app.register_blueprint(ai_bp)
app.register_blueprint(gmail_oauth_bp)
app.register_blueprint(recovery_bp)
app.register_blueprint(billing_bp)
app.register_blueprint(newsletter_bp)


# Inject branding/config variables into all templates for universal theming
@app.context_processor
def inject_branding():
    # Resolve brand/app names with DB settings taking precedence over config defaults
    # Prefer per-school settings, gracefully falling back to app/global config
    school_name = (
        get_setting("SCHOOL_NAME")
        or get_setting("APP_NAME")
        or app.config.get("APP_NAME")
    )
    brand = (
        get_setting("BRAND_NAME")
        or (school_name or "Fee Management System")
    )
    portal_title = (
        get_setting("PORTAL_TITLE")
        or app.config.get("PORTAL_TITLE")
        or "Fee Management portal"
    )
    app_title = ((school_name or brand) + f" {portal_title}").strip()
    # Resolve current academic term/year (best effort)
    cy, ct = None, None
    try:
        _conn = get_db_connection()
        try:
            cy, ct = get_or_seed_current_term(_conn)
        finally:
            _conn.close()
    except Exception:
        pass
    # Per-school first login marker for welcome banner
    first_login_at = None
    try:
        sid = session.get("school_id") if session else None
        if sid:
            _conn = get_db_connection()
            try:
                cur = _conn.cursor()
                cur.execute("SELECT first_login_at FROM schools WHERE id=%s", (sid,))
                row = cur.fetchone()
                if row is not None:
                    try:
                        first_login_at = row[0] if not isinstance(row, dict) else row.get("first_login_at")
                    except Exception:
                        first_login_at = None
            finally:
                _conn.close()
    except Exception:
        pass
    # Resolve per-school logo if uploaded
    school_logo_url = get_setting("SCHOOL_LOGO_URL")
    logo_primary = school_logo_url or app.config.get("LOGO_PRIMARY", "css/lovato_logo.jpg")
    logo_secondary = school_logo_url or app.config.get("LOGO_SECONDARY", "css/lovato_logo1.jpg")

    # Late payment penalty settings (per-school)
    late_kind = (get_setting("LATE_PENALTY_KIND") or "").strip()  # 'percent' or 'flat'
    try:
        late_value = float(get_setting("LATE_PENALTY_VALUE") or 0)
    except Exception:
        late_value = 0.0
    try:
        late_grace = int(float(get_setting("LATE_PENALTY_GRACE_DAYS") or 0))
    except Exception:
        late_grace = 0

    def _balance_class(value):
        try:
            val = float(value or 0)
        except Exception:
            return "balance-neutral"
        if val > 0:
            return "balance-negative"
        if val < 0:
            return "balance-positive"
        return "balance-neutral"

    return {
        "BRAND_NAME": brand,
        "PORTAL_TITLE": portal_title,
        "APP_TITLE": app_title,
        "LOGO_PRIMARY": logo_primary,
        "LOGO_SECONDARY": logo_secondary,
        "FAVICON": app.config.get("FAVICON", logo_primary),
        "BRAND_COLOR": app.config.get("BRAND_COLOR", "#059669"),
        "CURRENT_YEAR": cy,
        "CURRENT_TERM": ct,
        "PRO_PRICE_KES": app.config.get("PRO_PRICE_KES", 1500),
        # School profile (for printables/docs)
        "SCHOOL_NAME": school_name or brand,
        "SCHOOL_ADDRESS": get_setting("SCHOOL_ADDRESS") or "",
        "SCHOOL_PHONE": get_setting("SCHOOL_PHONE") or "",
        "SCHOOL_EMAIL": get_setting("SCHOOL_EMAIL") or "",
        "SCHOOL_WEBSITE": get_setting("SCHOOL_WEBSITE") or "",
        "SUPPORT_PHONE": app.config.get("SUPPORT_PHONE", "+254794656850"),
        "SCHOOL_FIRST_LOGIN_AT": first_login_at,
        # Penalties (for invoice rendering)
        "LATE_PENALTY_KIND": late_kind,
        "LATE_PENALTY_VALUE": late_value,
        "LATE_PENALTY_GRACE_DAYS": late_grace,
        # Raw uploaded logo path if present
        "SCHOOL_LOGO_URL": school_logo_url or "",
        "balance_class": _balance_class,
    }

# ---------- AUTH GUARD ----------
@app.before_request
def require_login_for_app():
    # Allow static files and login routes
    path = request.path or "/"
    allowed_prefixes = (
        "/static/",
        "/mpesa/callback",
        "/gmail/",      # OAuth start
        "/portal/",     # Student portal (token-based)
        "/g/",          # Guardian endpoints (public login)
        "/s/",          # Student password login
    )
    allowed_exact = {
        "/auth/login",
        "/auth/",
        "/auth/entry",
        "/auth/register",
        "/auth/register_school",
        "/admin/login",
        "/choose_school",
        "/healthz",
        "/readyz",
        "/livez",
        "/oauth2callback",  # OAuth redirect
        "/s/login",         # Student login
        "/g/login",         # Guardian login
        "/g/",              # Guardian index
        "/guardian_login",
        "/guardian_dashboard",
    }
    # Allow admin blueprint to self-guard; we just don't block its login
    if any(path.startswith(p) for p in allowed_prefixes) or path in allowed_exact:
        return None

    # If not logged in, redirect to login (except for admin routes which have own guard)
    if not session.get("user_logged_in"):
        # Let admin area be reachable separately (e.g., /admin, /admin/..)
        if path.startswith("/admin"):
            return None
        # API endpoints should also be protected
        if path not in ("/auth/", "/auth/entry", "/auth/login"):
            return redirect(url_for("auth.entry", next=path))
    # After login, require a selected school for app routes
    if not (path.startswith("/admin") or path.startswith("/auth")):
        if not session.get("school_id") and path != "/choose_school":
            return redirect(url_for("choose_school", next=path))
    return None


# ---------- HEALTH ENDPOINTS ----------
_APP_START_TS = datetime.now()


def _db_ping_ok() -> bool:
    try:
        conn = get_db_connection()
    except Exception:
        return False
    try:
        cur = conn.cursor()
        try:
            cur.execute("SELECT 1")
            _ = cur.fetchone()
            return True
        finally:
            try:
                cur.close()
            except Exception:
                pass
    finally:
        try:
            conn.close()
        except Exception:
            pass


@app.route("/healthz")
def healthz():
    """Basic liveness probe. Public and unauthenticated.

    Returns JSON with minimal info; does not require DB.
    """
    up_secs = max(0, int((datetime.now() - _APP_START_TS).total_seconds()))
    return jsonify({
        "ok": True,
        "status": "alive",
        "uptime_seconds": up_secs,
        "version": app.config.get("APP_NAME", "Fee Management System"),
    })


@app.route("/readyz")
def readyz():
    """Readiness probe. Checks DB connectivity best-effort."""
    db_ok = _db_ping_ok()
    code = 200 if db_ok else 503
    return jsonify({"ok": db_ok, "db": db_ok}), code


@app.route("/livez")
def livez():
    return jsonify({"ok": True})


# Convenience: /login -> /auth/login
@app.route("/login")
def login_redirect():
    return redirect(url_for("auth.entry"))


# Convenience: /guardian -> /g/login
@app.route("/guardian")
@app.route("/guardian/")
def guardian_redirect():
    return redirect(url_for("guardian.guardian_login"))

# Public aliases for guardian endpoints (UX-friendly names)
@app.route("/guardian_login", methods=["GET", "POST"])
def guardian_login_alias():
    try:
        # Delegate to blueprint route
        return app.view_functions.get("guardian.guardian_login")()
    except Exception:
        return redirect(url_for("guardian.guardian_login"))

@app.route("/guardian_dashboard", methods=["GET"])
def guardian_dashboard_alias():
    try:
        return app.view_functions.get("guardian.guardian_dashboard")()
    except Exception:
        return redirect(url_for("guardian.guardian_dashboard"))

@app.route("/make_payment", methods=["POST"])
def make_payment_alias():
    try:
        return app.view_functions.get("guardian.guardian_make_payment")()
    except Exception:
        # If alias fails, return 400 style response without crashing
        from flask import jsonify
        return jsonify({"ok": False, "error": "Payment endpoint unavailable"}), 400


# ---------- DATABASE CONNECTION ----------
def get_db_connection():
    """Establish a connection to the MySQL database."""
    # Prefer credentials from SQLALCHEMY_DATABASE_URI to avoid hardcoding
    uri = app.config.get("SQLALCHEMY_DATABASE_URI", "")
    host = os.environ.get("DB_HOST", "localhost")
    user = os.environ.get("DB_USER", "root")
    password = os.environ.get("DB_PASSWORD", "")
    database = os.environ.get("DB_NAME", "school_fee_db")

    if uri:
        try:
            parsed = urlparse(uri)
            # Only attempt to parse for MySQL-style URIs
            if parsed.scheme.startswith("mysql"):
                if parsed.hostname:
                    host = parsed.hostname
                if parsed.username:
                    user = parsed.username
                if parsed.password:
                    password = parsed.password
                if parsed.path and len(parsed.path) > 1:
                    database = parsed.path.lstrip("/")
        except Exception:
            # Fall back to env/defaults if parsing fails
            pass

    # Optional MySQL TLS settings via environment (off by default for local dev)
    kwargs = dict(host=host, user=user, password=password, database=database)
    try:
        ssl_disabled = os.environ.get("DB_SSL_DISABLED", "0").strip().lower() in ("1", "true", "yes")
        require_tls = os.environ.get("DB_SSL_REQUIRE", "0").strip().lower() in ("1", "true", "yes")
        ssl_ca = os.environ.get("DB_SSL_CA", "").strip() or None
        ssl_cert = os.environ.get("DB_SSL_CERT", "").strip() or None
        ssl_key = os.environ.get("DB_SSL_KEY", "").strip() or None
        ssl_verify = os.environ.get("DB_SSL_VERIFY", "1").strip().lower() not in ("0", "false", "no")

        if ssl_disabled:
            kwargs["ssl_disabled"] = True  # type: ignore[assignment]
        elif ssl_ca or require_tls:
            # Only pass SSL args when explicitly requested or CA provided
            if ssl_ca:
                kwargs["ssl_ca"] = ssl_ca  # type: ignore[assignment]
            if ssl_cert:
                kwargs["ssl_cert"] = ssl_cert  # type: ignore[assignment]
            if ssl_key:
                kwargs["ssl_key"] = ssl_key  # type: ignore[assignment]
            kwargs["ssl_verify_cert"] = ssl_verify  # type: ignore[assignment]
        # Else: do not set any SSL options -> plaintext connection (e.g., localhost)
    except Exception:
        # Fall back to non-TLS if env parsing fails
        pass

    return mysql.connector.connect(**kwargs)


def get_db_connection_readonly():
    """Get a short-lived, read-only connection for analytics (non-blocking).

    - autocommit enabled so SELECTs do not hold open transactions
    - isolation lowered to READ COMMITTED to avoid unnecessary locks
    - marks session READ ONLY where supported
    """
    conn = get_db_connection()
    try:
        try:
            conn.autocommit = True  # type: ignore[attr-defined]
        except Exception:
            pass
        try:
            cur = conn.cursor()
            try:
                cur.execute("SET SESSION TRANSACTION ISOLATION LEVEL READ COMMITTED")
            except Exception:
                pass
            try:
                cur.execute("SET SESSION TRANSACTION READ ONLY")
            except Exception:
                pass
            try:
                cur.close()
            except Exception:
                pass
        except Exception:
            pass
    except Exception:
        pass
    return conn


def _bootstrap_db_safely():
    """Attempt DB schema bootstrap without blocking app startup."""
    try:
        db = get_db_connection()
    except Exception:
        return
    try:
        # Ensure multi-tenant scaffolding exists
        ensure_schools_table(db)
        ensure_school_id_columns(
            db,
            (
                "students",
                "payments",
                "credit_operations",
                "credit_transfers",
                "academic_terms",
                "student_enrollments",
                "term_fees",
            ),
        )
        ensure_students_credit_column(db)
        ensure_credit_ops_table(db)
        ensure_credit_transfers_table(db)
        # Academic term scaffolding
        ensure_academic_terms_table(db)
        ensure_payments_term_columns(db)
        # User tables for multi-user (premium-ready)
        try:
            ensure_user_tables(db)
        except Exception:
            pass
        ensure_student_enrollments_table(db)
        ensure_term_fees_table(db)
        # Strengthen per-school uniqueness where safe
        try:
            ensure_unique_indices_per_school(db)
        except Exception:
            pass
        # Ensure newsletters storage
        try:
            ensure_newsletters_table(db)
        except Exception:
            pass
        # Audit feature removed: no audit table initialization
    except Exception:
        pass
    finally:
        try:
            db.close()
        except Exception:
            pass


_BOOTSTRAP_DONE = False

@app.before_request
def _run_bootstrap_once():
    global _BOOTSTRAP_DONE
    if not _BOOTSTRAP_DONE:
        _BOOTSTRAP_DONE = True
        _bootstrap_db_safely()

# ---------- SCHOOL SELECTION ----------
@app.route("/choose_school", methods=["GET", "POST"])
def choose_school():
    if request.method == "POST":
        raw_code = (request.form.get("school_code") or "").strip()
        name = (request.form.get("school_name") or "").strip()
        code = slugify_code(raw_code or name)
        if not code:
            flash("Enter a school name or code.", "warning")
            return redirect(url_for("choose_school"))
        db = get_db_connection()
        created = False
        try:
            # Check if exists first to decide on bootstrapping
            cur = db.cursor()
            cur.execute("SELECT id FROM schools WHERE code=%s", (code,))
            existing = cur.fetchone()
            sid = None
            if existing:
                sid = int(existing[0]) if not isinstance(existing, dict) else int(existing.get("id"))
            if not sid:
                # Restrict creating new schools to admin area
                if not session.get("is_admin"):
                    db.close()
                    flash("Only an administrator can create a new school. Please sign in to Admin and create it there.", "warning")
                    return redirect(url_for("admin.login", next=url_for("choose_school")))
                sid = get_or_create_school(db, code=code, name=name or code)
                created = True
            if created:
                try:
                    bootstrap_new_school(db, sid, name or code, code)
                except Exception:
                    pass
        finally:
            db.close()
        session["school_id"] = sid
        session["school_code"] = code
        if created:
            flash("School created. Sign in with default credentials: user / 9133.", "info")
        next_url = request.args.get("next") or request.form.get("next") or url_for("dashboard")
        return redirect(next_url)
    return render_template("choose_school.html", next_url=request.args.get("next", ""))
# ---------- DASHBOARD ----------
@app.route("/")
def dashboard():
    """Main dashboard with summary cards and recent payments."""
    db = get_db_connection()
    cursor = db.cursor(dictionary=True)

    # Resolve current academic context
    try:
        current_year, current_term = get_or_seed_current_term(db)
    except Exception:
        current_year, current_term = None, None

    # Totals
    cursor.execute("SELECT COUNT(*) AS total FROM students WHERE school_id=%s", (session.get("school_id"),))
    total_students = cursor.fetchone()["total"]

    # Total collected for current term (fallback to lifetime if context missing)
    if current_year and current_term in (1, 2, 3):
        cursor.execute(
            "SELECT COALESCE(SUM(amount), 0) AS total_collected FROM payments WHERE method <> 'Credit Transfer' AND school_id=%s AND year=%s AND term=%s",
            (session.get("school_id"), current_year, current_term),
        )
    else:
        cursor.execute(
            "SELECT COALESCE(SUM(amount), 0) AS total_collected FROM payments WHERE method <> 'Credit Transfer' AND school_id=%s",
            (session.get("school_id"),),
        )
    total_collected = cursor.fetchone()["total_collected"]

    # Detect correct balance column
    cursor.execute("SHOW COLUMNS FROM students LIKE 'balance'")
    has_balance = bool(cursor.fetchone())
    column = "balance" if has_balance else "fee_balance"

    cursor.execute(f"SELECT COALESCE(SUM({column}), 0) AS total_balance FROM students WHERE school_id=%s", (session.get("school_id"),))
    total_balance = cursor.fetchone()["total_balance"]

    # Total credit
    cursor.execute("SELECT COALESCE(SUM(credit), 0) AS total_credit FROM students WHERE school_id=%s", (session.get("school_id"),))
    total_credit = cursor.fetchone()["total_credit"]

    # Recent payments
    recent_sql_base = (
        """
        SELECT p.id, s.name, s.class_name, p.amount, p.method, p.date
        FROM payments p
        JOIN students s ON p.student_id = s.id
        WHERE p.school_id=%s {extra}
        ORDER BY p.date DESC
        LIMIT 5
        """
    )
    if current_year and current_term in (1, 2, 3):
        cursor.execute(
            recent_sql_base.format(extra="AND p.year=%s AND p.term=%s"),
            (session.get("school_id"), current_year, current_term),
        )
    else:
        cursor.execute(
            recent_sql_base.format(extra=""),
            (session.get("school_id"),),
        )
    recent_payments = cursor.fetchall()

    # Recently added students (last 10)
    cursor.execute(
        """
        SELECT id, name, admission_no, class_name
        FROM students
        WHERE school_id=%s
        ORDER BY id DESC
        LIMIT 10
        """,
        (session.get("school_id"),),
    )
    recent_students = cursor.fetchall()

    # Compute this term outstanding (fees for current term minus payments in current term)
    term_outstanding = 0.0
    try:
        if current_year and current_term in (1, 2, 3):
            school_id = session.get("school_id")
            cursor.execute("SELECT id FROM students WHERE school_id=%s", (school_id,))
            _rows = cursor.fetchall() or []
            ids = [r.get("id") for r in _rows]
            if ids:
                def _in_clause(seq):
                    return ",".join(["%s"] * len(seq))
                # Itemized per student
                cursor.execute(
                    f"SELECT student_id, COALESCE(SUM(amount),0) AS tsum FROM student_term_fee_items WHERE year=%s AND term=%s AND student_id IN ({_in_clause(ids)}) GROUP BY student_id",
                    (current_year, current_term, *ids),
                )
                items_map = {r["student_id"]: float(r.get("tsum") or 0) for r in (cursor.fetchall() or [])}
                # Legacy flat per student
                cursor.execute(
                    f"SELECT student_id, COALESCE(SUM(fee_amount),0) AS tsum FROM term_fees WHERE year=%s AND term=%s AND student_id IN ({_in_clause(ids)}) GROUP BY student_id",
                    (current_year, current_term, *ids),
                )
                legacy_map = {r["student_id"]: float(r.get("tsum") or 0) for r in (cursor.fetchall() or [])}
                total_fee = sum(items_map.values())
                for sid in ids:
                    if sid not in items_map:
                        total_fee += float(legacy_map.get(sid) or 0)
                # Payments this term
                cursor.execute(
                    "SELECT COALESCE(SUM(amount),0) AS t FROM payments WHERE method <> 'Credit Transfer' AND school_id=%s AND year=%s AND term=%s",
                    (school_id, current_year, current_term),
                )
                total_pay = float((cursor.fetchone() or {}).get("t", 0) or 0)
                term_outstanding = max(total_fee - total_pay, 0.0)
    except Exception:
        term_outstanding = 0.0

    db.close()
    return render_template(
        "dashboard.html",
        total_students=total_students,
        total_fees_collected=total_collected,
        pending_balance=total_balance,
        total_credit=total_credit,
        recent_payments=recent_payments,
        recent_students=recent_students,
        term_outstanding=term_outstanding
    )


# ---------- REAL-TIME DASHBOARD API ----------
@app.route("/api/dashboard_data")
def dashboard_data():
    """Return real-time dashboard totals."""
    db = get_db_connection()
    cursor = db.cursor(dictionary=True)

    cursor.execute("SHOW COLUMNS FROM students LIKE 'balance'")
    has_balance = bool(cursor.fetchone())
    column = "balance" if has_balance else "fee_balance"

    cursor.execute("SELECT COUNT(*) AS total_students FROM students WHERE school_id=%s", (session.get("school_id"),))
    total_students = cursor.fetchone()["total_students"]

    # Use current term for totals where applicable
    try:
        cy, ct = get_or_seed_current_term(db)
    except Exception:
        cy, ct = None, None
    if cy and ct in (1, 2, 3):
        cursor.execute(
            "SELECT COALESCE(SUM(amount), 0) AS total_collected FROM payments WHERE method <> 'Credit Transfer' AND school_id=%s AND year=%s AND term=%s",
            (session.get("school_id"), cy, ct),
        )
    else:
        cursor.execute(
            "SELECT COALESCE(SUM(amount), 0) AS total_collected FROM payments WHERE method <> 'Credit Transfer' AND school_id=%s",
            (session.get("school_id"),),
        )
    total_collected = cursor.fetchone()["total_collected"]

    cursor.execute(f"SELECT COALESCE(SUM({column}), 0) AS total_balance FROM students WHERE school_id=%s", (session.get("school_id"),))
    total_balance = cursor.fetchone()["total_balance"]

    cursor.execute("SELECT COALESCE(SUM(credit), 0) AS total_credit FROM students WHERE school_id=%s", (session.get("school_id"),))
    total_credit = cursor.fetchone()["total_credit"]

    # Compute term outstanding if we have a current year/term
    term_outstanding = 0.0
    expected_term_total = 0.0
    term_collection_rate = 0.0
    try:
        if cy and ct in (1, 2, 3):
            school_id = session.get("school_id")
            cursor.execute("SELECT id FROM students WHERE school_id=%s", (school_id,))
            _rows = cursor.fetchall() or []
            ids = [r.get("id") for r in _rows]
            if ids:
                def _in_clause(seq):
                    return ",".join(["%s"] * len(seq))
                cursor.execute(
                    f"SELECT student_id, COALESCE(SUM(amount),0) AS tsum FROM student_term_fee_items WHERE year=%s AND term=%s AND student_id IN ({_in_clause(ids)}) GROUP BY student_id",
                    (cy, ct, *ids),
                )
                items_map = {r["student_id"]: float(r.get("tsum") or 0) for r in (cursor.fetchall() or [])}
                cursor.execute(
                    f"SELECT student_id, COALESCE(SUM(fee_amount),0) AS tsum FROM term_fees WHERE year=%s AND term=%s AND student_id IN ({_in_clause(ids)}) GROUP BY student_id",
                    (cy, ct, *ids),
                )
                legacy_map = {r["student_id"]: float(r.get("tsum") or 0) for r in (cursor.fetchall() or [])}
                # Build expected map per student, preferring itemized records; fall back to legacy where missing
                expected_map = dict(items_map)
                for sid in ids:
                    if sid not in expected_map:
                        expected_map[sid] = float(legacy_map.get(sid) or 0)
                expected_term_total = float(sum(expected_map.values()) or 0)

                # Payments per student for current term
                cursor.execute(
                    f"SELECT student_id, COALESCE(SUM(amount),0) AS tsum FROM payments WHERE method <> 'Credit Transfer' AND school_id=%s AND year=%s AND term=%s AND student_id IN ({_in_clause(ids)}) GROUP BY student_id",
                    (school_id, cy, ct, *ids),
                )
                paid_map = {r["student_id"]: float(r.get("tsum") or 0) for r in (cursor.fetchall() or [])}

                # Outstanding computed per-student (overpay on one student does not offset another)
                term_outstanding = 0.0
                paid_capped_sum = 0.0
                for sid in ids:
                    exp = float(expected_map.get(sid) or 0)
                    paid = float(paid_map.get(sid) or 0)
                    term_outstanding += max(exp - paid, 0.0)
                    paid_capped_sum += min(paid, exp)

                # Collection rate per-student sum, excluding overpay
                if expected_term_total > 0:
                    term_collection_rate = round((paid_capped_sum / expected_term_total) * 100.0, 1)
    except Exception:
        term_outstanding = 0.0
        expected_term_total = 0.0
        term_collection_rate = 0.0

    db.close()
    resp = jsonify({
        "total_students": total_students,
        "total_collected": float(total_collected or 0),
        "total_balance": float(total_balance or 0),
        "total_credit": float(total_credit or 0),
        "term_outstanding": float(term_outstanding or 0),
        "expected_term_total": float(expected_term_total or 0),
        "term_collection_rate": float(term_collection_rate or 0)
    })
    resp.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
    resp.headers["Pragma"] = "no-cache"
    resp.headers["Expires"] = "0"
    return resp


# ---------- LEDGER VIEW ----------
@app.route("/students/<int:student_id>/ledger")
def student_ledger(student_id: int):
    db = get_db_connection()
    try:
        ensure_ledger_table(db)
        cur = db.cursor(dictionary=True)
        cur.execute(
            "SELECT * FROM ledger_entries WHERE school_id=%s AND student_id=%s ORDER BY ts DESC, id DESC",
            (session.get("school_id"), student_id),
        )
        entries = cur.fetchall() or []
    finally:
        db.close()
    return render_template("student_ledger.html", entries=entries, student_id=student_id)


# ---------- FORECAST API ----------
@app.route("/api/forecast_collections")
def forecast_collections():
    db = get_db_connection()
    cur = db.cursor(dictionary=True)
    try:
        cur.execute(
            """
            SELECT DATE_FORMAT(date, '%Y-%m') AS ym, SUM(amount) AS total
            FROM payments
            WHERE method <> 'Credit Transfer' AND school_id=%s
            GROUP BY DATE_FORMAT(date, '%Y-%m')
            ORDER BY ym DESC
            LIMIT 6
            """,
            (session.get("school_id"),),
        )
        rows = cur.fetchall() or []
        hist = [float(r.get('total') or 0) for r in rows][::-1]
        avg = sum(hist) / len(hist) if hist else 0.0
        forecast = [round(avg, 2)] * 3
        return jsonify({"ok": True, "method": "moving_average", "horizon": 3, "forecast": forecast})
    finally:
        db.close()


# ---------- STUDENTS ----------
@app.route("/students")
def students():
    db = get_db_connection()
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT * FROM students WHERE school_id=%s ORDER BY id DESC", (session.get("school_id"),))
    students = cursor.fetchall()
    db.close()
    return render_template("students.html", students=students)


# ---------- GLOBAL SEARCH API ----------
@app.route("/api/search")
def global_search():
    q = (request.args.get("q") or "").strip()
    if not q:
        return jsonify({"students": [], "payments": []})
    db = get_db_connection_readonly()
    cur = db.cursor(dictionary=True)
    try:
        # Determine balance column
        cur.execute("SHOW COLUMNS FROM students LIKE 'balance'")
        has_balance = bool(cur.fetchone())
        bal_col = "balance" if has_balance else "fee_balance"

        like = f"%{q}%"
        # Students by name/admission/class
        cur.execute(
            f"""
            SELECT id, name, class_name, admission_no, COALESCE({bal_col},0) AS balance, COALESCE(credit,0) AS credit
            FROM students
            WHERE school_id=%s AND (name LIKE %s OR admission_no LIKE %s OR class_name LIKE %s)
            ORDER BY name ASC
            LIMIT 15
            """,
            (session.get("school_id"), like, like, like),
        )
        students = cur.fetchall() or []

        # Payments: search by reference, student name, and optionally by exact amount if q is numeric
        amt = None
        try:
            amt = float(q.replace(",", ""))
        except Exception:
            amt = None
        if amt is not None:
            cur.execute(
                """
                SELECT p.id, p.reference, p.amount, DATE_FORMAT(p.date, '%Y-%m-%d') AS date, s.name AS student_name
                FROM payments p
                LEFT JOIN students s ON s.id = p.student_id
                WHERE p.school_id=%s AND (
                      p.reference LIKE %s OR s.name LIKE %s OR p.amount = %s
                )
                ORDER BY p.date DESC
                LIMIT 10
                """,
                (session.get("school_id"), like, like, amt),
            )
        else:
            cur.execute(
                """
                SELECT p.id, p.reference, p.amount, DATE_FORMAT(p.date, '%Y-%m-%d') AS date, s.name AS student_name
                FROM payments p
                LEFT JOIN students s ON s.id = p.student_id
                WHERE p.school_id=%s AND (
                      p.reference LIKE %s OR s.name LIKE %s
                )
                ORDER BY p.date DESC
                LIMIT 10
                """,
                (session.get("school_id"), like, like),
            )
        payments = cur.fetchall() or []
    finally:
        db.close()
    return jsonify({"students": students, "payments": payments})


@app.route("/add_student", methods=["GET", "POST"])
def add_student():
    """Add a new student (reject duplicates by admission number only)."""
    if request.method == "POST":
        name = request.form["name"].strip()
        admission_no = request.form.get("admission_no", "").strip()
        class_name = request.form["class_name"].strip()
        phone = request.form.get("phone", "").strip()
        # Email for reminders; save to whichever column exists (email or parent_email)
        email_val = (request.form.get("email") or "").strip()
        total_fees = float(request.form.get("total_fees", 0))

        db = get_db_connection()
        cursor = db.cursor(dictionary=True)

        # Detect correct column
        cursor.execute("SHOW COLUMNS FROM students LIKE 'balance'")
        has_balance = bool(cursor.fetchone())
        cursor.execute("SHOW COLUMNS FROM students LIKE 'fee_balance'")
        has_fee_balance = bool(cursor.fetchone())
        # Detect optional phone column
        cursor.execute("SHOW COLUMNS FROM students LIKE 'phone'")
        has_phone_col = bool(cursor.fetchone())
        # Detect possible email columns
        cursor.execute("SHOW COLUMNS FROM students LIKE 'email'")
        has_email_col = bool(cursor.fetchone())
        cursor.execute("SHOW COLUMNS FROM students LIKE 'parent_email'")
        has_parent_email_col = bool(cursor.fetchone())

        # --- DUPLICATE CHECK (Admission No only) ---
        if admission_no:
            cursor.execute(
                """
                SELECT id FROM students
                WHERE LOWER(admission_no) = LOWER(%s) AND school_id=%s
                """,
                (admission_no, session.get("school_id")),
            )
            existing = cursor.fetchone()
            if existing:
                db.close()
                flash(f"Admission Number '{admission_no}' already exists in the system.", "warning")
                return redirect(url_for("students"))

        # Insert (build dynamically to include optional phone/email columns)
        if not (has_balance or has_fee_balance):
            db.close()
            flash("No valid balance column found in 'students' table!", "error")
            return redirect(url_for("students"))

        cols = ["name", "admission_no", "class_name"]
        params_list = [name, admission_no or None, class_name]
        if has_phone_col:
            cols.append("phone")
            params_list.append(phone or None)
        # Prefer 'email' if present; else 'parent_email' if present
        email_col = None
        if has_email_col:
            email_col = "email"
        elif has_parent_email_col:
            email_col = "parent_email"
        if email_col:
            cols.append(email_col)
            params_list.append(email_val or None)
        # Balance column
        if has_balance:
            cols.append("balance")
        else:
            cols.append("fee_balance")
        params_list.append(total_fees)
        # Credit and school
        cols.append("credit")
        cols.append("school_id")
        params_list.append(0)
        params_list.append(session.get("school_id"))

        placeholders = ", ".join(["%s"] * len(params_list))
        sql = f"INSERT INTO students ({', '.join(cols)}) VALUES ({placeholders})"
        params = tuple(params_list)

        try:
            cursor.execute(sql, params)
            student_id = cursor.lastrowid
            try:
                ensure_student_enrollments_table(db)
                cy, _ct = get_or_seed_current_term(db)
                cur2 = db.cursor()
                cur2.execute(
                    "INSERT IGNORE INTO student_enrollments (student_id, year, class_name, opening_balance, status, school_id) VALUES (%s,%s,%s,%s,%s,%s)",
                    (student_id, cy, class_name, total_fees, "active", session.get("school_id")),
                )
            except Exception:
                pass
            # Seed default student portal password as the admission number
            try:
                if admission_no:
                    ensure_student_portal_columns(db)
                    hp = hash_password(admission_no)
                    cur3 = db.cursor()
                    cur3.execute("UPDATE students SET portal_password_hash=%s WHERE id=%s", (hp, student_id))
            except Exception:
                pass
            db.commit()
            # audit removed
            flash(f"Student '{name}' added successfully!", "success")
        except Exception as e:
            db.rollback()
            flash(f"Error adding student: {e}", "error")
        finally:
            db.close()

        return redirect(url_for("students"))

    return render_template("add_student.html")


# ---------- IMPORT STUDENTS (CSV) ----------
@app.route("/import_students", methods=["GET", "POST"])
def import_students():
    """Bulk import students from a CSV file.

    Expected columns (header row recommended):
      - name, admission_no, class_name, phone, email, total_fees

    Rules:
      - Duplicate check uses admission_no within the same school (if provided).
      - total_fees maps to opening balance for the current year enrollment and to student balance/fee_balance.
      - Phone and email columns are optional in DB; they are used if columns exist.
    """
    if request.method == "GET":
        return render_template("import_students.html")

    file = request.files.get("file")
    if not file or not file.filename:
        flash("Please choose a CSV file to upload.", "warning")
        return redirect(url_for("import_students"))

    filename = secure_filename(file.filename)
    try:
        raw = file.stream.read().decode("utf-8-sig")
    except Exception:
        try:
            raw = file.stream.read().decode("latin-1")
        except Exception:
            flash("Could not read CSV file. Ensure it is UTF-8 encoded.", "error")
            return redirect(url_for("import_students"))

    # Parse CSV
    f = StringIO(raw)
    reader = csv.DictReader(f)

    # If no header present, fall back to simple reader and map columns by index
    used_dict_reader = True
    if not reader.fieldnames or len([c for c in reader.fieldnames if c]) <= 1:
        used_dict_reader = False
        f.seek(0)
        reader = csv.reader(f)

    db = get_db_connection()
    cur = db.cursor(dictionary=True)

    # Detect schema columns
    cur.execute("SHOW COLUMNS FROM students LIKE 'balance'")
    has_balance = bool(cur.fetchone())
    cur.execute("SHOW COLUMNS FROM students LIKE 'fee_balance'")
    has_fee_balance = bool(cur.fetchone())
    cur.execute("SHOW COLUMNS FROM students LIKE 'phone'")
    has_phone_col = bool(cur.fetchone())
    cur.execute("SHOW COLUMNS FROM students LIKE 'email'")
    has_email_col = bool(cur.fetchone())
    cur.execute("SHOW COLUMNS FROM students LIKE 'parent_email'")
    has_parent_email_col = bool(cur.fetchone())

    if not (has_balance or has_fee_balance):
        db.close()
        flash("No valid balance column found in 'students' table!", "error")
        return redirect(url_for("students"))

    imported, duplicates, errors = 0, 0, 0
    detail_errors = []

    try:
        ensure_student_enrollments_table(db)
        cy, _ct = get_or_seed_current_term(db)
    except Exception:
        cy = None

    try:
        for row in reader:
            try:
                if used_dict_reader:
                    name = (row.get("name") or "").strip()
                    admission_no = (row.get("admission_no") or row.get("admission") or "").strip()
                    class_name = (row.get("class_name") or row.get("class") or "").strip()
                    phone = (row.get("phone") or "").strip()
                    email_val = (row.get("email") or row.get("parent_email") or "").strip()
                    tf_raw = (row.get("total_fees") or row.get("fees") or row.get("balance") or "0").strip()
                else:
                    # Positional mapping: [name, admission_no, class_name, phone, email, total_fees]
                    cols = list(row)
                    name = (cols[0] if len(cols) > 0 else "").strip()
                    admission_no = (cols[1] if len(cols) > 1 else "").strip()
                    class_name = (cols[2] if len(cols) > 2 else "").strip()
                    phone = (cols[3] if len(cols) > 3 else "").strip()
                    email_val = (cols[4] if len(cols) > 4 else "").strip()
                    tf_raw = (cols[5] if len(cols) > 5 else "0").strip()

                if not name or not class_name:
                    errors += 1
                    detail_errors.append("Missing required name or class_name")
                    continue

                try:
                    total_fees = float(tf_raw.replace(",", "")) if tf_raw else 0.0
                except Exception:
                    total_fees = 0.0

                # Normalize phone if util available
                try:
                    phone = normalize_phone(phone) if phone else phone
                except Exception:
                    pass

                # Duplicate check by admission_no (if provided)
                if admission_no:
                    cur.execute(
                        "SELECT id FROM students WHERE LOWER(admission_no)=LOWER(%s) AND school_id=%s",
                        (admission_no, session.get("school_id")),
                    )
                    if cur.fetchone():
                        duplicates += 1
                        continue

                # Build insert
                cols = ["name", "admission_no", "class_name"]
                params = [name, admission_no or None, class_name]
                if has_phone_col:
                    cols.append("phone")
                    params.append(phone or None)
                email_col = None
                if has_email_col:
                    email_col = "email"
                elif has_parent_email_col:
                    email_col = "parent_email"
                if email_col:
                    cols.append(email_col)
                    params.append(email_val or None)
                # Balance column
                cols.append("balance" if has_balance else "fee_balance")
                params.append(total_fees)
                # Credit and school id
                cols += ["credit", "school_id"]
                params += [0, session.get("school_id")]

                placeholders = ", ".join(["%s"] * len(params))
                sql = f"INSERT INTO students ({', '.join(cols)}) VALUES ({placeholders})"
                cur.execute(sql, tuple(params))
                sid = cur.lastrowid

                # Enrollment with opening balance
                try:
                    if cy is not None:
                        cur2 = db.cursor()
                        cur2.execute(
                            "INSERT IGNORE INTO student_enrollments (student_id, year, class_name, opening_balance, status, school_id) VALUES (%s,%s,%s,%s,%s,%s)",
                            (sid, cy, class_name, total_fees, "active", session.get("school_id")),
                        )
                except Exception:
                    pass

                imported += 1
            except Exception as e:
                errors += 1
                detail_errors.append(str(e))

        db.commit()
    except Exception as e:
        db.rollback()
        db.close()
        flash(f"Import failed: {e}", "error")
        return redirect(url_for("import_students"))
    finally:
        try:
            db.close()
        except Exception:
            pass

    # Summary
    msg = f"Imported {imported} student(s)."
    if duplicates:
        msg += f" Skipped {duplicates} duplicate(s)."
    if errors:
        msg += f" {errors} error(s) occurred."
    flash(msg, "success" if imported and not errors else ("warning" if imported else "error"))
    return redirect(url_for("students"))


# ---------- EDIT STUDENT ----------
@app.route("/student/<int:student_id>/edit", methods=["GET", "POST"])
def edit_student(student_id):
    """Edit existing student details and update balance/fee_balance accordingly."""
    db = get_db_connection()
    cursor = db.cursor(dictionary=True)

    # Detect schema
    cursor.execute("SHOW COLUMNS FROM students LIKE 'balance'")
    has_balance = bool(cursor.fetchone())
    cursor.execute("SHOW COLUMNS FROM students LIKE 'fee_balance'")
    has_fee_balance = bool(cursor.fetchone())
    cursor.execute("SHOW COLUMNS FROM students LIKE 'phone'")
    has_phone_col = bool(cursor.fetchone())
    # Optional: one-time retention flag
    try:
        cursor.execute("SHOW COLUMNS FROM students LIKE 'retain_next_year'")
        has_retain = bool(cursor.fetchone())
    except Exception:
        has_retain = False

    cursor.execute("SELECT * FROM students WHERE id = %s AND school_id=%s", (student_id, session.get("school_id")))
    student = cursor.fetchone()

    if not student:
        db.close()
        flash("Student not found.", "error")
        return redirect(url_for("students"))

    if request.method == "POST":
        name = request.form.get("name", student.get("name", "")).strip()
        admission_no = request.form.get("admission_no", student.get("admission_no", "")).strip()
        class_name = request.form.get("class_name", student.get("class_name", "")).strip()
        balance_val = float(request.form.get("balance", student.get("balance") or student.get("fee_balance") or 0))
        phone_val = (request.form.get("phone") or student.get("phone") or "").strip()
        email_val = (request.form.get("email") or student.get("email") or student.get("parent_email") or "").strip()

        # Enforce unique admission_no if changed and provided
        if admission_no and admission_no.lower() != (student.get("admission_no") or "").lower():
            cursor.execute(
                "SELECT id FROM students WHERE LOWER(admission_no) = LOWER(%s) AND school_id=%s",
                (admission_no, session.get("school_id"))
            )
            exists = cursor.fetchone()
            if exists:
                db.close()
                flash("Admission Number already exists.", "warning")
                return redirect(url_for("edit_student", student_id=student_id))

        # Detect optional email columns for update
        cursor.execute("SHOW COLUMNS FROM students LIKE 'email'")
        has_email_col = bool(cursor.fetchone())
        cursor.execute("SHOW COLUMNS FROM students LIKE 'parent_email'")
        has_parent_email_col = bool(cursor.fetchone())

        # Build dynamic update
        sets = ["name = %s", "admission_no = %s", "class_name = %s"]
        params = [name, admission_no or None, class_name]
        if has_balance:
            sets.append("balance = %s")
            params.append(balance_val)
        elif has_fee_balance:
            sets.append("fee_balance = %s")
            params.append(balance_val)
        if has_phone_col:
            sets.append("phone = %s")
            params.append(phone_val or None)
        # Update the appropriate email column if present
        if has_email_col:
            sets.append("email = %s")
            params.append(email_val or None)
        elif has_parent_email_col:
            sets.append("parent_email = %s")
            params.append(email_val or None)
        # Retention toggle
        if has_retain:
            retain_flag = 1 if (request.form.get("retain_next_year") == "on") else 0
            sets.append("retain_next_year = %s")
            params.append(retain_flag)
        params.append(student_id)
        params.append(session.get("school_id"))

        try:
            cursor.execute(f"UPDATE students SET {', '.join(sets)} WHERE id = %s AND school_id=%s", tuple(params))
            db.commit()
            # audit removed
            flash("Student updated successfully!", "success")
        except Exception as e:
            db.rollback()
            flash(f"Error updating student: {e}", "error")
        finally:
            db.close()
        return redirect(url_for("student_detail", student_id=student_id))

    # Ensure template has a 'balance' key for display regardless of schema
    if "balance" not in student:
        student["balance"] = student.get("fee_balance")

    db.close()
    return render_template("edit_student.html", student=student)


@app.route("/delete_student/<int:student_id>", methods=["POST"])
def delete_student(student_id):
    """Delete student and related payments."""
    db = get_db_connection()
    cursor = db.cursor()
    cursor.execute("DELETE FROM payments WHERE student_id = %s AND school_id=%s", (student_id, session.get("school_id")))
    cursor.execute("DELETE FROM students WHERE id = %s AND school_id=%s", (student_id, session.get("school_id")))
    db.commit()
    # audit removed
    db.close()
    flash("Student deleted successfully!", "success")
    return redirect(url_for("students"))


# ---------- SEARCH ----------
@app.route("/search_student")
def search_student():
    """Search students for the current school.

    Works across different profile schemas by coalescing balance columns
    and searching common fields (name, class, admission, phone/email when present).
    """
    query = (request.args.get("query") or "").strip()
    db = get_db_connection_readonly()
    cursor = db.cursor(dictionary=True)

    sid = session.get("school_id")
    if not sid:
        return jsonify([])

    if query:
        like = f"%{query}%"
        # Detect optional columns so we don't reference missing fields
        cursor.execute("SHOW COLUMNS FROM students LIKE 'phone'")
        has_phone = bool(cursor.fetchone())
        cursor.execute("SHOW COLUMNS FROM students LIKE 'email'")
        has_email = bool(cursor.fetchone())
        cursor.execute("SHOW COLUMNS FROM students LIKE 'parent_email'")
        has_parent_email = bool(cursor.fetchone())

        where_parts = [
            "name LIKE %s",
            "class_name LIKE %s",
            "admission_no LIKE %s",
        ]
        params = [sid, like, like, like]
        if has_phone:
            where_parts.append("COALESCE(phone,'') LIKE %s")
            params.append(like)
        if has_email or has_parent_email:
            email_expr = []
            if has_email:
                email_expr.append("COALESCE(email,'') LIKE %s")
                params.append(like)
            if has_parent_email:
                email_expr.append("COALESCE(parent_email,'') LIKE %s")
                params.append(like)
            where_parts.append("(" + " OR ".join(email_expr) + ")")

        where_clause = " OR ".join(where_parts)
        sql = (
            f"SELECT id, name, class_name, admission_no, "
            f"COALESCE(balance, fee_balance, 0) AS balance, "
            f"COALESCE(credit, 0) AS credit "
            f"FROM students "
            f"WHERE school_id=%s AND ({where_clause}) "
            f"ORDER BY name ASC"
        )
        cursor.execute(sql, tuple(params))
    else:
        cursor.execute(
            """
            SELECT id, name, class_name, admission_no,
                   COALESCE(balance, fee_balance, 0) AS balance,
                   COALESCE(credit, 0) AS credit
            FROM students
            WHERE school_id=%s
            ORDER BY id DESC
            """,
            (sid,),
        )

    students = cursor.fetchall()
    db.close()
    return jsonify(students)


# ---------- DUPLICATE CHECK API ----------
@app.route("/check_student_exists")
def check_student_exists():
    """AJAX check if Admission Number already exists."""
    admission_no = request.args.get("admission_no", "").strip().lower()
    if not admission_no:
        return jsonify({"exists": False})

    db = get_db_connection()
    cursor = db.cursor(dictionary=True)
    cursor.execute(
        """
        SELECT id FROM students
        WHERE LOWER(admission_no) = %s AND school_id=%s
        """,
        (admission_no, session.get("school_id")),
    )
    exists = bool(cursor.fetchone())
    db.close()

    return jsonify({"exists": exists})


# ---------- EXPORT REPORTS ----------
@app.route("/export_students")
def export_students():
    """Export all student records as CSV."""
    db = get_db_connection()
    cursor = db.cursor(dictionary=True)
    cursor.execute(
        """
        SELECT 
            name AS 'Name', 
            admission_no AS 'Admission No', 
            class_name AS 'Class', 
            COALESCE(balance, fee_balance) AS 'Balance (KES)', 
            COALESCE(credit, 0) AS 'Credit (KES)'
        FROM students
        WHERE school_id=%s
        ORDER BY class_name, name
        """,
        (session.get("school_id"),),
    )
    students = cursor.fetchall()
    db.close()

    if not students:
        return Response("No students found.", mimetype="text/plain")

    output = StringIO()
    writer = csv.DictWriter(output, fieldnames=students[0].keys())
    writer.writeheader()
    writer.writerows(students)
    output.seek(0)

    return Response(
        output.getvalue(),
        mimetype="text/csv",
        headers={"Content-Disposition": "attachment; filename=students_report.csv"}
    )


@app.route("/export_payments")
def export_payments():
    """Export payment records as CSV, with optional timeframe filters.

    Accepts optional query params:
      - year: numeric year
      - term: 1/2/3
      - start_date, end_date: ISO dates to bound by p.date (inclusive)
    If both year and term are provided, they take precedence over date range.
    """
    db = get_db_connection()
    cur = db.cursor(dictionary=True)
    sid = session.get("school_id")

    year = request.args.get("year")
    term = request.args.get("term")
    start_date = request.args.get("start_date")
    end_date = request.args.get("end_date")
    try:
        y_val = int(year) if (year or "").isdigit() else None
    except Exception:
        y_val = None
    try:
        t_val = int(term) if (term or "").isdigit() else None
    except Exception:
        t_val = None

    sql = (
        "SELECT s.name AS 'Student Name', s.admission_no AS 'Admission No', s.class_name AS 'Class', "
        "p.year AS 'Year', p.term AS 'Term', p.amount AS 'Amount (KES)', p.method AS 'Method', "
        "p.reference AS 'Reference', p.date AS 'Date' "
        "FROM payments p JOIN students s ON s.id = p.student_id WHERE p.method <> 'Credit Transfer' AND p.school_id=%s"
    )
    params = [sid]
    if y_val and t_val:
        sql += " AND p.year=%s AND p.term=%s"; params += [y_val, t_val]
    else:
        if y_val:
            sql += " AND p.year=%s"; params.append(y_val)
        if start_date:
            sql += " AND p.date >= %s"; params.append(start_date)
        if end_date:
            sql += " AND p.date <= %s"; params.append(end_date)
    sql += " ORDER BY p.date DESC"
    cur.execute(sql, tuple(params))
    payments = cur.fetchall()
    db.close()
    output = StringIO()
    fieldnames = payments[0].keys() if payments else ["Student Name", "Admission No", "Class", "Year", "Term", "Amount (KES)", "Method", "Reference", "Date"]
    writer = csv.DictWriter(output, fieldnames=fieldnames)
    writer.writeheader()
    writer.writerows(payments)
    output.seek(0)

    return Response(
        output.getvalue(),
        mimetype="text/csv",
        headers={"Content-Disposition": "attachment; filename=payments_report.csv"}
    )


@app.route("/export_analytics_csv")
def export_analytics_csv():
    """Analytics exports disabled."""
    return Response("Not Found", status=404)

    ds = (request.args.get("dataset") or "").strip().lower()
    sid = session.get("school_id")
    if not sid:
        return Response("Missing school context", status=400)

    db = get_db_connection()
    cur = db.cursor(dictionary=True)

    # Helper to get current term
    try:
        cy, ct = get_or_seed_current_term(db)
    except Exception:
        cy, ct = None, None

    try:
        if ds == "expected_by_class":
            # Compute expected and capped collected per class (current term)
            cur.execute("SELECT id, class_name FROM students WHERE school_id=%s", (sid,))
            stu = cur.fetchall() or []
            id2class = {r["id"]: (r.get("class_name") or "") for r in stu}
            ids = list(id2class.keys())
            expected_by_class = {}
            paid_capped_by_class = {}
            if ids and cy and ct in (1, 2, 3):
                def _in_clause(seq):
                    return ",".join(["%s"] * len(seq))
                cur.execute(
                    f"SELECT student_id, COALESCE(SUM(amount),0) AS tsum FROM student_term_fee_items WHERE year=%s AND term=%s AND student_id IN ({_in_clause(ids)}) GROUP BY student_id",
                    (cy, ct, *ids),
                )
                items_map = {r["student_id"]: float(r.get("tsum") or 0) for r in (cur.fetchall() or [])}
                cur.execute(
                    f"SELECT student_id, COALESCE(SUM(fee_amount),0) AS tsum FROM term_fees WHERE year=%s AND term=%s AND student_id IN ({_in_clause(ids)}) GROUP BY student_id",
                    (cy, ct, *ids),
                )
                legacy_map = {r["student_id"]: float(r.get("tsum") or 0) for r in (cur.fetchall() or [])}
                expected_map = dict(items_map)
                for sid_i in ids:
                    if sid_i not in expected_map:
                        expected_map[sid_i] = float(legacy_map.get(sid_i) or 0)
                cur.execute(
                    f"SELECT student_id, COALESCE(SUM(amount),0) AS tsum FROM payments WHERE method <> 'Credit Transfer' AND school_id=%s AND year=%s AND term=%s AND student_id IN ({_in_clause(ids)}) GROUP BY student_id",
                    (sid, cy, ct, *ids),
                )
                paid_map = {r["student_id"]: float(r.get("tsum") or 0) for r in (cur.fetchall() or [])}
                for sid_i in ids:
                    cls = id2class.get(sid_i, "")
                    exp = float(expected_map.get(sid_i) or 0)
                    paidv = float(paid_map.get(sid_i) or 0)
                    expected_by_class[cls] = expected_by_class.get(cls, 0.0) + exp
                    paid_capped_by_class[cls] = paid_capped_by_class.get(cls, 0.0) + min(paidv, exp)
            rows = []
            for cls in sorted(set(list(expected_by_class.keys()) + list(paid_capped_by_class.keys()))):
                exp = float(expected_by_class.get(cls, 0.0))
                col = float(paid_capped_by_class.get(cls, 0.0))
                rows.append({"Class": cls or "", "Expected (KES)": exp, "Collected (Capped) (KES)": col, "Gap (KES)": max(exp - col, 0.0)})
            out = StringIO(); fn = ["Class", "Expected (KES)", "Collected (Capped) (KES)", "Gap (KES)"]
            w = csv.DictWriter(out, fieldnames=fn); w.writeheader(); [w.writerow(r) for r in rows]; out.seek(0)
            return Response(out.getvalue(), mimetype="text/csv", headers={"Content-Disposition": "attachment; filename=expected_by_class.csv"})

        elif ds == "aging":
            # Aging by last payment (students with positive balance)
            bal_col = "balance"
            try:
                cur.execute("SHOW COLUMNS FROM students LIKE 'balance'")
                if not cur.fetchone():
                    bal_col = "fee_balance"
            except Exception:
                bal_col = "fee_balance"
            cur.execute(
                f"""
                SELECT s.id, s.name, s.class_name, COALESCE({bal_col},0) AS bal,
                       (SELECT MAX(p.date) FROM payments p WHERE p.student_id=s.id AND p.school_id=%s) AS last_pay
                FROM students s
                WHERE s.school_id=%s
                ORDER BY s.class_name, s.name
                """,
                (sid, sid),
            )
            from datetime import date as _date
            today = _date.today()
            def _bucket(days):
                return "0-30" if days <= 30 else ("31-60" if days <= 60 else ("61-90" if days <= 90 else ">90"))
            out = StringIO(); fn = ["Student", "Class", "Balance (KES)", "Days Since Last Payment", "Bucket"]
            w = csv.DictWriter(out, fieldnames=fn); w.writeheader()
            for r in (cur.fetchall() or []):
                bal = float(r.get("bal") or 0)
                if bal <= 0: continue
                lp = r.get("last_pay"); days = 9999
                try:
                    if lp is not None:
                        d = lp.date() if hasattr(lp, "date") else lp
                        days = max(0, (today - d).days)
                except Exception:
                    pass
                w.writerow({
                    "Student": r.get("name") or "",
                    "Class": r.get("class_name") or "",
                    "Balance (KES)": bal,
                    "Days Since Last Payment": days,
                    "Bucket": _bucket(days),
                })
            out.seek(0)
            return Response(out.getvalue(), mimetype="text/csv", headers={"Content-Disposition": "attachment; filename=aging_last_payment.csv"})

        elif ds == "method_trend":
            cur.execute(
                """
                SELECT DATE_FORMAT(date, '%Y-%m') AS Month, method AS Method, COALESCE(SUM(amount),0) AS Total
                FROM payments
                WHERE method <> 'Credit Transfer' AND school_id=%s AND date >= DATE_SUB(CURRENT_DATE, INTERVAL 6 MONTH)
                GROUP BY DATE_FORMAT(date, '%Y-%m'), method
                ORDER BY Month, Method
                """,
                (sid,),
            )
            rows = cur.fetchall() or []
            out = StringIO(); fn = ["Month", "Method", "Total"]
            w = csv.DictWriter(out, fieldnames=fn); w.writeheader(); [w.writerow(r) for r in rows]; out.seek(0)
            return Response(out.getvalue(), mimetype="text/csv", headers={"Content-Disposition": "attachment; filename=method_trend_6m.csv"})

        elif ds == "weekday_trend":
            cur.execute(
                """
                SELECT DAYOFWEEK(date) AS dow, COALESCE(SUM(amount),0) AS total
                FROM payments
                WHERE method <> 'Credit Transfer' AND school_id=%s AND date >= (CURRENT_DATE - INTERVAL 182 DAY)
                GROUP BY DAYOFWEEK(date)
                ORDER BY dow
                """,
                (sid,),
            )
            rows = cur.fetchall() or []
            # Map to names
            name = {1:"Sun",2:"Mon",3:"Tue",4:"Wed",5:"Thu",6:"Fri",7:"Sat"}
            out = StringIO(); fn = ["Weekday", "Total"]
            w = csv.DictWriter(out, fieldnames=fn); w.writeheader()
            for r in rows:
                w.writerow({"Weekday": name.get(int(r.get("dow") or 0), str(r.get("dow"))), "Total": float(r.get("total") or 0)})
            out.seek(0)
            return Response(out.getvalue(), mimetype="text/csv", headers={"Content-Disposition": "attachment; filename=weekday_trend.csv"})

        elif ds == "forecast":
            cur.execute(
                """
                SELECT DATE_FORMAT(date, '%Y-%m') AS Month, COALESCE(SUM(amount),0) AS Total
                FROM payments
                WHERE method <> 'Credit Transfer' AND school_id=%s AND date >= DATE_SUB(CURRENT_DATE, INTERVAL 12 MONTH)
                GROUP BY DATE_FORMAT(date, '%Y-%m')
                ORDER BY Month
                """,
                (sid,),
            )
            rows = cur.fetchall() or []
            last_vals = [float(r.get("Total") or 0) for r in rows][-3:]
            next_val = float(sum(last_vals)/len(last_vals)) if last_vals else 0.0
            out = StringIO(); fn = ["Month", "Total"]
            w = csv.DictWriter(out, fieldnames=fn); w.writeheader(); [w.writerow(r) for r in rows]
            w.writerow({"Month": "Forecast Next", "Total": next_val}); out.seek(0)
            return Response(out.getvalue(), mimetype="text/csv", headers={"Content-Disposition": "attachment; filename=forecast_next.csv"})

        elif ds == "due_aging":
            # Aggregate outstanding by due date buckets for current term invoices
            try:
                cy, ct = get_or_seed_current_term(db)
            except Exception:
                cy, ct = None, None
            if not (cy and ct in (1, 2, 3)):
                return Response("No current term", status=400)
            cur.execute(
                """
                SELECT i.student_id, i.due_date, i.total
                FROM invoices i JOIN students s ON s.id = i.student_id
                WHERE i.year=%s AND i.term=%s AND s.school_id=%s
                """,
                (cy, ct, sid),
            )
            inv = cur.fetchall() or []
            cur.execute(
                """
                SELECT student_id, COALESCE(SUM(amount),0) AS total
                FROM payments
                WHERE method <> 'Credit Transfer' AND school_id=%s AND year=%s AND term=%s
                GROUP BY student_id
                """,
                (sid, cy, ct),
            )
            paid_map = {r["student_id"]: float(r.get("total") or 0) for r in (cur.fetchall() or [])}
            from datetime import date as _date
            today = _date.today()
            buckets = {"Not Due":0.0, "0-30":0.0, "31-60":0.0, "61-90":0.0, ">90":0.0}
            for r in inv:
                total = float(r.get("total") or 0)
                paid = float(paid_map.get(r.get("student_id"), 0.0))
                out_amt = max(total - min(paid, total), 0.0)
                d = r.get("due_date")
                if out_amt <= 0 or d is None:
                    continue
                try:
                    days = (today - d).days
                except Exception:
                    days = 0
                if days <= 0:
                    buckets["Not Due"] += out_amt
                elif days <= 30:
                    buckets["0-30"] += out_amt
                elif days <= 60:
                    buckets["31-60"] += out_amt
                elif days <= 90:
                    buckets["61-90"] += out_amt
                else:
                    buckets[">90"] += out_amt
            out = StringIO(); fn=["Bucket","Outstanding (KES)"]
            w=csv.DictWriter(out, fieldnames=fn); w.writeheader()
            for k in ["Not Due","0-30","31-60","61-90",">90"]:
                w.writerow({"Bucket": k, "Outstanding (KES)": buckets[k]})
            out.seek(0)
            return Response(out.getvalue(), mimetype="text/csv", headers={"Content-Disposition": "attachment; filename=aging_due_date.csv"})

        elif ds == "timeliness":
            # Per-student on-time vs late paid for current term
            try:
                cy, ct = get_or_seed_current_term(db)
            except Exception:
                cy, ct = None, None
            if not (cy and ct in (1, 2, 3)):
                return Response("No current term", status=400)
            cur.execute(
                """
                SELECT i.student_id, i.total, i.due_date, s.name, s.class_name
                FROM invoices i JOIN students s ON s.id = i.student_id
                WHERE i.year=%s AND i.term=%s AND s.school_id=%s
                """,
                (cy, ct, sid),
            )
            inv = cur.fetchall() or []
            cur.execute(
                """
                SELECT p.student_id,
                       COALESCE(SUM(CASE WHEN p.date <= i.due_date THEN p.amount ELSE 0 END),0) AS paid_on_time,
                       COALESCE(SUM(CASE WHEN p.date > i.due_date THEN p.amount ELSE 0 END),0) AS paid_late
                FROM payments p JOIN invoices i
                  ON i.student_id = p.student_id AND i.year=p.year AND i.term=p.term
                WHERE p.method <> 'Credit Transfer' AND p.school_id=%s AND p.year=%s AND p.term=%s
                GROUP BY p.student_id
                """,
                (sid, cy, ct),
            )
            tmap = {r["student_id"]: {"on": float(r.get("paid_on_time") or 0), "late": float(r.get("paid_late") or 0)} for r in (cur.fetchall() or [])}
            out = StringIO(); fn=["Student","Class","Expected (KES)","Paid On Time (KES)","Paid Late (KES)"]
            w=csv.DictWriter(out, fieldnames=fn); w.writeheader()
            for r in inv:
                sid_i = r.get("student_id")
                total = float(r.get("total") or 0)
                mt = tmap.get(sid_i, {"on":0.0, "late":0.0})
                on_cap = min(mt["on"], total)
                tot_paid = min(mt["on"] + mt["late"], total)
                late_cap = max(tot_paid - on_cap, 0.0)
                w.writerow({
                    "Student": r.get("name") or "",
                    "Class": r.get("class_name") or "",
                    "Expected (KES)": total,
                    "Paid On Time (KES)": on_cap,
                    "Paid Late (KES)": late_cap,
                })
            out.seek(0)
            return Response(out.getvalue(), mimetype="text/csv", headers={"Content-Disposition": "attachment; filename=payment_timeliness.csv"})

        else:
            return Response("Unknown dataset", status=400)
    finally:
        try:
            db.close()
        except Exception:
            pass


# ---------- STUDENT DETAIL ----------
@app.route("/student/<int:student_id>")
def student_detail(student_id):
    """View student profile and payments."""
    db = get_db_connection()
    cursor = db.cursor(dictionary=True)

    cursor.execute("SELECT * FROM students WHERE id = %s AND school_id=%s", (student_id, session.get("school_id")))
    student = cursor.fetchone()

    cursor.execute(
        """
        SELECT * FROM payments
        WHERE student_id = %s AND school_id=%s
        ORDER BY date DESC
        """,
        (student_id, session.get("school_id"))
    )
    payments = cursor.fetchall()
    # Overpay records from credit operations (if table exists)
    try:
        ensure_credit_ops_table(db)
        cursor.execute(
            """
            SELECT ts, amount, reference, method
            FROM credit_operations
            WHERE student_id = %s AND op_type = 'overpay' AND school_id=%s
            ORDER BY ts DESC
            """,
            (student_id, session.get("school_id"))
        )
        overpays = cursor.fetchall()
        # Credit transfer records (both directions) if table exists
        ensure_credit_transfers_table(db)
        # Outgoing transfers with destination name
        cursor.execute(
            """
            SELECT ct.*, s.name AS to_name
            FROM credit_transfers ct
            JOIN students s ON s.id = ct.to_student_id
            WHERE ct.from_student_id = %s AND ct.school_id=%s
            ORDER BY ct.ts DESC
            """,
            (student_id, session.get("school_id"))
        )
        transfers_out = cursor.fetchall()
        # Incoming transfers with source name
        cursor.execute(
            """
            SELECT ct.*, s.name AS from_name
            FROM credit_transfers ct
            JOIN students s ON s.id = ct.from_student_id
            WHERE ct.to_student_id = %s AND ct.school_id=%s
            ORDER BY ct.ts DESC
            """,
            (student_id, session.get("school_id"))
        )
        transfers_in = cursor.fetchall()
    except Exception:
        overpays = []
        transfers_out = []
        transfers_in = []
    db.close()

    if not student:
        flash("Student not found.", "error")
        return redirect(url_for("students"))

    # Audit: record profile view with context
    try:
        log_event("view_student_profile", target=f"student:{student_id}", detail=f"{student.get("name")}")
    except Exception:
        pass

    # âœ… Pass datetime to template
    return render_template(
        "view_student.html",
        student=student,
        payments=payments,
        overpays=overpays,
        transfers_out=transfers_out,
        transfers_in=transfers_in,
        datetime=datetime,
    )


# ---------- PAYMENTS ----------
@app.route("/payments", methods=["GET", "POST"])
def payments():
    """Add or list payments."""
    db = get_db_connection()
    cursor = db.cursor(dictionary=True)

    if request.method == "POST":
        student_id = request.form["student_id"]
        amount = float(request.form["amount"])
        method = request.form["method"]
        reference = (request.form.get("reference") or "").strip()
        payment_date = datetime.now()
        carry_overpay = (request.form.get("carry_overpay") in ("1", "true", "on", "yes"))
        # Academic context (Term/Year)
        try:
            form_year = request.form.get("year", type=int)
            form_term = request.form.get("term", type=int)
        except Exception:
            form_year, form_term = None, None

        # Enforce term state: only allow payments in OPEN term
        try:
            cy_check, ct_check = get_or_seed_current_term(db)
            cur_gate = db.cursor()
            sid_gate = session.get("school_id")
            cur_gate.execute(
                "SELECT status FROM academic_terms WHERE year=%s AND term=%s AND school_id=%s",
                (cy_check, ct_check, sid_gate),
            )
            trow = cur_gate.fetchone()
            if trow is not None:
                status_val = trow[0] if not isinstance(trow, dict) else (trow.get("status") or "DRAFT")
                if status_val != "OPEN":
                    db.close()
                    flash("Payments are locked until the current term is OPEN.", "warning")
                    return redirect(url_for("payments"))
        except Exception:
            pass

        # Detect correct balance column
        cursor.execute("SHOW COLUMNS FROM students LIKE 'balance'")
        has_balance = bool(cursor.fetchone())
        column = "balance" if has_balance else "fee_balance"

        cursor.execute(f"SELECT {column}, credit FROM students WHERE id = %s AND school_id=%s", (student_id, session.get("school_id")))
        student = cursor.fetchone()

        # Fetch student's contact info for notifications (email-first)
        cursor.execute("SHOW COLUMNS FROM students LIKE 'phone'")
        _has_phone_col = bool(cursor.fetchone())
        # Detect email column preference: 'email' then 'parent_email'
        _email_col = None
        cursor.execute("SHOW COLUMNS FROM students LIKE 'email'")
        if bool(cursor.fetchone()):
            _email_col = 'email'
        else:
            cursor.execute("SHOW COLUMNS FROM students LIKE 'parent_email'")
            if bool(cursor.fetchone()):
                _email_col = 'parent_email'

        student_name = None
        student_phone = None
        student_email = None
        # Retrieve core name and available contact columns
        select_cols = ["name", "class_name"]
        if _has_phone_col:
            select_cols.append("phone")
        if _email_col:
            select_cols.append(_email_col)
        cols_sql = ", ".join(select_cols)
        cursor.execute(f"SELECT {cols_sql} FROM students WHERE id = %s AND school_id=%s", (student_id, session.get("school_id")))
        row = cursor.fetchone() or {}
        student_name = row.get("name")
        if _has_phone_col:
            student_phone = row.get("phone")
        if _email_col:
            student_email = row.get(_email_col)

        if not student:
            db.close()
            flash("Student not found!", "error")
            return redirect(url_for("payments"))

        current_balance = float(student[column] or 0)
        current_credit = float(student["credit"] or 0)

        if amount > current_balance:
            overpaid = amount - current_balance
            new_balance = 0
            # if admin chooses to carry to next term, do NOT add to credit now
            new_credit = current_credit if carry_overpay else (current_credit + overpaid)
        else:
            new_balance = current_balance - amount
            new_credit = current_credit
        # Track total overpay for logging (guaranteed non-negative)
        overpaid_total = max(amount - current_balance, 0.0)

        # Enforce unique reference if provided (case-insensitive)
        if reference:
            cursor.execute(
                "SELECT COUNT(*) AS cnt FROM payments WHERE LOWER(reference) = LOWER(%s) AND school_id=%s",
                (reference, session.get("school_id")),
            )
            ref_exists = (cursor.fetchone() or {}).get("cnt", 0)
            if ref_exists:
                db.close()
                flash("Reference already exists. Use a unique REF.", "warning")
                return redirect(url_for("payments", student_id=student_id))

        # Ensure term/year columns exist and compute defaults
        try:
            ensure_payments_term_columns(db)
        except Exception:
            pass
        if not (form_year and form_term in (1, 2, 3)):
            try:
                cy, ct = get_or_seed_current_term(db)
            except Exception:
                cy, ct = payment_date.year, None
        else:
            cy, ct = form_year, form_term

        cursor.execute(
            """
            INSERT INTO payments (student_id, amount, method, term, year, reference, date, school_id)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
            """,
            (student_id, amount, method, ct, cy, reference or None, payment_date, session.get("school_id")),
        )
        payment_id = cursor.lastrowid

        cursor.execute(
            f"UPDATE students SET {column} = %s, credit = %s WHERE id = %s AND school_id=%s",
            (new_balance, new_credit, student_id, session.get("school_id")),
        )
        db.commit()
        # Ledger only
        try:
            ensure_ledger_table(db)
            add_entry(
                db,
                school_id=int(session.get("school_id")),
                student_id=int(student_id),
                entry_type='credit',
                amount=float(amount),
                ref=reference or None,
                description=f"Payment via {method}",
                link_type='payment',
                link_id=int(payment_id),
            )
        except Exception:
            pass
        try:
            log_event(
                "record_payment",
                target=f"payment:{payment_id}",
                detail=f"Paid KES {amount:,.2f} for student {student_name or student.get('name')}",
            )
        except Exception:
            pass

        # Optionally carry any overpay to next term by creating a forward payment entry
        carried_payment_id = None
        if carry_overpay and overpaid_total > 0 and ct in (1, 2, 3):
            try:
                # work out next (year, term)
                next_term = 1 if ct == 3 else (ct + 1)
                next_year = (cy + 1) if ct == 3 else cy
                carry_ref = (reference or f"PMT-{payment_id}") + "-CF"
                cursor.execute(
                    """
                    INSERT INTO payments (student_id, amount, method, term, year, reference, date, school_id)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                    """,
                    (student_id, overpaid_total, "Carry Forward", next_term, next_year, carry_ref, payment_date, session.get("school_id")),
                )
                carried_payment_id = cursor.lastrowid
                db.commit()
                # Ledger entry for carry-forward credit
                try:
                    ensure_ledger_table(db)
                    add_entry(
                        db,
                        school_id=int(session.get("school_id")),
                        student_id=int(student_id),
                        entry_type='credit',
                        amount=float(overpaid_total),
                        ref=carry_ref,
                        description=f"Carry forward from {cy} T{ct}",
                        link_type='payment',
                        link_id=int(carried_payment_id),
                    )
                except Exception:
                    pass
            except Exception:
                # If carry write fails, fall back to adding to credit to avoid loss
                try:
                    cursor.execute(
                        "UPDATE students SET credit = credit + %s WHERE id = %s AND school_id=%s",
                        (overpaid_total, student_id, session.get("school_id")),
                    )
                    db.commit()
                except Exception:
                    pass

        # If there was an overpay, record it in credit operations for student detail visibility
        try:
            if overpaid_total > 0:
                ensure_credit_ops_table(db)
                cur2 = db.cursor()
                cur2.execute(
                    "INSERT INTO credit_operations (ts, actor, student_id, op_type, amount, reference, method, meta, school_id) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s)",
                    (
                        datetime.utcnow(),
                        session.get("username"),
                        int(student_id),
                        "overpay",
                        overpaid_total,
                        reference or None,
                        method,
                        json.dumps({
                            "payment_id": int(payment_id),
                            "carried": bool(carry_overpay),
                            "carried_payment_id": int(carried_payment_id) if carried_payment_id else None
                        }),
                        session.get("school_id"),
                    ),
                )
                db.commit()
        except Exception:
            # Do not block normal flow if logging fails
            pass
        # audit removed

        # Send a modern HTML receipt via Email when available
        brand = (
            (get_setting("SCHOOL_NAME") or "").strip()
            or (get_setting("BRAND_NAME") or "").strip()
            or (app.config.get("APP_NAME") or "Your School")
        )
        pretty_ref = reference or "N/A"
        email_subject = f"Payment receipt for {student_name} - KES {amount:,.2f}"
        email_sent = False
        if student_email:
            # Build receipt link for verification and quick access
            try:
                receipt_url = url_for('payment_receipt', payment_id=payment_id, _external=True)
            except Exception:
                receipt_url = ""
            # Render an email-friendly HTML receipt
            try:
                email_html = render_template(
                    "email_receipt.html",
                    brand=brand,
                    student_name=student_name,
                    class_name=row.get("class_name") if isinstance(row, dict) else None,
                    amount=amount,
                    method=method,
                    reference=pretty_ref,
                    new_balance=new_balance,
                    new_credit=new_credit,
                    receipt_url=receipt_url,
                    payment_id=payment_id,
                    year=cy,
                    term=ct,
                )
            except Exception:
                # Graceful fallback to a plain text summary
                email_html = None
            try:
                # Ensure a clean, user-friendly subject line (ASCII-only for safety)
                email_subject = f"Payment receipt for {student_name} - KES {amount:,.2f}"
                if email_html:
                    email_sent = gmail_send_email_html(student_email, email_subject, email_html)
                else:
                    email_body = (
                        f"{brand}: Payment received.\n"
                        f"Hi {student_name}, thank you for your payment of KES {amount:,.2f} via {method} (Ref: {pretty_ref}).\n"
                        f"Your new balance is KES {new_balance:,.2f}. Credit on account: KES {new_credit:,.2f}.\n"
                        + (f"View receipt: {receipt_url}" if receipt_url else "")
                    )
                    email_sent = gmail_send_email(student_email, email_subject, email_body)
            except Exception:
                email_sent = False
            if not email_sent:
                # Fallback to Flask-Mail (SMTP) only if configured
                try:
                    smtp_server = (app.config.get('MAIL_SERVER') or '').strip()
                    smtp_user = (app.config.get('MAIL_USERNAME') or '').strip()
                    smtp_pass = (app.config.get('MAIL_PASSWORD') or '').strip()
                    if smtp_server and smtp_user and smtp_pass:
                        from flask_mail import Message
                        from extensions import mail
                        sender = (
                            app.config.get('MAIL_SENDER')
                            or app.config.get('MAIL_DEFAULT_SENDER')
                            or (get_setting('SCHOOL_EMAIL') or None)
                            or app.config.get('MAIL_USERNAME')
                            or None
                        )
                        if email_html:
                            m = Message(subject=email_subject, sender=sender, recipients=[student_email], html=email_html)
                        else:
                            m = Message(subject=email_subject, sender=sender, recipients=[student_email], body=email_body)
                        mail.send(m)
                        email_sent = True
                    else:
                        email_sent = False
                except Exception:
                    email_sent = False
        if not student_email:
            flash("Payment recorded. Add an email to send receipts.", "info")
        elif not email_sent:
            flash("Payment recorded, but sending email receipt failed.", "warning")
        db.close()

        flash(f"Payment of KES {amount:,.2f} recorded! Remaining balance: KES {new_balance:,.2f}, Credit: KES {new_credit:,.2f}", "success")
        return redirect(url_for("payments"))

    # âœ… NEW: read preselected student from query param
    selected_student_id = request.args.get("student_id", type=int)

    cursor.execute(
        """
        SELECT p.*, s.name AS student_name, s.class_name
        FROM payments p
        JOIN students s ON p.student_id = s.id
        WHERE p.school_id=%s
        ORDER BY p.date DESC
        """,
        (session.get("school_id"),),
    )
    payments = cursor.fetchall()
    
    cursor.execute("SELECT id, name FROM students WHERE school_id=%s ORDER BY name ASC", (session.get("school_id"),))
    students = cursor.fetchall()
    db.close()

    return render_template("payments.html", payments=payments, students=students, selected_student_id=selected_student_id)


@app.route("/payments/<int:payment_id>/delete", methods=["POST"])
def delete_payment(payment_id: int):
    db = get_db_connection()
    cursor = db.cursor(dictionary=True)
    cursor.execute(
        "SELECT p.id, p.student_id, p.amount, p.method, s.name AS student_name FROM payments p JOIN students s ON s.id=p.student_id WHERE p.id=%s AND p.school_id=%s",
        (payment_id, session.get("school_id")),
    )
    record = cursor.fetchone()
    if not record:
        db.close()
        flash("Payment not found or already deleted.", "warning")
        return redirect(url_for("payments"))
    cursor.execute("DELETE FROM payments WHERE id=%s AND school_id=%s", (payment_id, session.get("school_id")))
    db.commit()
    try:
        log_event(
            "delete_payment",
            target=f"payment:{payment_id}",
            detail=f"Removed payment KES {record.get('amount') or 0:,.2f} for {record.get('student_name')}",
        )
    except Exception:
        pass
    flash("Payment deleted.", "info")
    db.close()
    return redirect(url_for("payments"))


# ---------- PAYMENT RECEIPT (Printable) ----------
@app.route("/payments/<int:payment_id>/receipt")
def payment_receipt(payment_id: int):
    """Render a compact, one-page printable receipt for a payment."""
    db = get_db_connection()
    cursor = db.cursor(dictionary=True)

    # Fetch payment with student info
    cursor.execute(
        """
        SELECT p.*, s.name AS student_name, s.class_name, s.id AS sid
        FROM payments p
        JOIN students s ON s.id = p.student_id
        WHERE p.id = %s AND p.school_id=%s
        """,
        (payment_id, session.get("school_id")),
    )
    payment = cursor.fetchone()
    if not payment:
        db.close()
        flash("Payment not found.", "error")
        return redirect(url_for("payments"))

    # Determine current balance/credit (for display only)
    try:
        cursor.execute("SHOW COLUMNS FROM students LIKE 'balance'")
        has_balance = bool(cursor.fetchone())
        column = "balance" if has_balance else "fee_balance"

        cursor.execute(
            f"SELECT {column} AS balance, credit FROM students WHERE id = %s AND school_id=%s",
            (payment.get("sid"), session.get("school_id")),
        )
        sc = cursor.fetchone() or {}
        current_balance = float((sc.get("balance") or 0))
        current_credit = float((sc.get("credit") or 0))
    except Exception:
        current_balance = None
        current_credit = None

    db.close()

    # Prefer school-specific branding for receipt header
    brand = (
        (get_setting("SCHOOL_NAME") or "").strip()
        or (get_setting("BRAND_NAME") or "").strip()
        or (app.config.get("APP_NAME") or "").strip()
        or (app.config.get("BRAND_NAME") or "").strip()
        or "School"
    )
    # Build a verification URL for authenticity (scannable)
    try:
        verify_url = url_for('payment_receipt', payment_id=payment_id, _external=True)
    except Exception:
        verify_url = ""

    # Build signed QR payload with key details for authenticity
    try:
        pdate = payment.get("date")
        if hasattr(pdate, "strftime"):
            date_str = pdate.strftime("%Y-%m-%d %H:%M")
        else:
            date_str = str(pdate or "")
        qr_payload = {
            "t": "receipt",
            "rid": int(payment.get("id")),
            "sid": int(payment.get("sid")),
            "amt": round(float(payment.get("amount") or 0.0), 2),
            "cur": "KES",
            "name": payment.get("student_name") or "",
            "cls": payment.get("class_name") or "",
            "dt": date_str,
            "m": payment.get("method") or "",
            "ref": payment.get("reference") or "",
            "term": payment.get("term") or "",
            "year": payment.get("year") or "",
            "school_id": session.get("school_id"),
        }
        canon = json.dumps(qr_payload, sort_keys=True, separators=(",", ":")).encode("utf-8")
        sig = hmac.new(app.secret_key.encode("utf-8"), canon, hashlib.sha256).hexdigest()[:20]
        qr_payload["sig"] = sig
        auth_qr_data = json.dumps(qr_payload, separators=(",", ":"))
    except Exception:
        auth_qr_data = ""

    return render_template(
        "receipt.html",
        brand=brand,
        payment=payment,
        current_balance=current_balance,
        current_credit=current_credit,
        payment_link=app.config.get("PAYMENT_LINK", ""),
        verify_url=verify_url,
        auth_qr_data=auth_qr_data,
    )


# ---------- PAYMENT RECEIPT (PDF Download) ----------
@app.route("/payments/<int:payment_id>/receipt.pdf")
def payment_receipt_pdf(payment_id: int):
    """Generate a PDF receipt for a payment with school details."""
    db = get_db_connection()
    cursor = db.cursor(dictionary=True)

    cursor.execute(
        """
        SELECT p.id, p.student_id AS sid, s.name AS student_name, s.class_name, p.amount, p.method,
               p.reference, p.date, p.term, p.year
        FROM payments p
        JOIN students s ON s.id = p.student_id
        WHERE p.id = %s AND p.school_id=%s
        """,
        (payment_id, session.get("school_id")),
    )
    payment = cursor.fetchone()
    if not payment:
        db.close()
        flash("Payment not found.", "error")
        return redirect(url_for("payments"))

    cursor.execute("SHOW COLUMNS FROM students LIKE 'balance'")
    has_balance = bool(cursor.fetchone())
    bal_col = "balance" if has_balance else "fee_balance"
    cursor.execute(
        f"SELECT COALESCE({bal_col}, 0) AS bal, COALESCE(credit, 0) AS credit FROM students WHERE id=%s AND school_id=%s",
        (payment.get("sid"), session.get("school_id")),
    )
    srow = cursor.fetchone() or {"bal": 0.0, "credit": 0.0}
    db.close()

    school_name = (
        get_setting("SCHOOL_NAME")
        or get_setting("APP_NAME")
        or app.config.get("APP_NAME")
        or app.config.get("BRAND_NAME")
        or get_setting("BRAND_NAME")
        or "School"
    )
    school_address = get_setting("SCHOOL_ADDRESS") or ""
    school_phone = get_setting("SCHOOL_PHONE") or ""
    school_email = get_setting("SCHOOL_EMAIL") or ""
    school_website = get_setting("SCHOOL_WEBSITE") or ""

    from io import BytesIO
    from reportlab.lib.pagesizes import A5
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.pdfgen import canvas

    # Styled, modern PDF layout to mirror HTML receipt
    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=A5)
    width, height = A5
    x_margin = 14 * mm
    content_gap = 6 * mm

    # Brand colors (match HTML theme)
    brand_indigo = colors.HexColor("#4338ca")
    brand_cyan = colors.HexColor("#06b6d4")
    light_bg = colors.HexColor("#eef2ff")
    soft_border = colors.HexColor("#e2e8f0")

    # Header bar
    header_h = 28 * mm
    c.setFillColor(brand_indigo)
    c.setStrokeColor(brand_indigo)
    c.rect(0, height - header_h, width, header_h, fill=1, stroke=0)

    # Optional logo on header
    logo_w = 16 * mm
    logo_h = 16 * mm
    logo_x = x_margin
    logo_y = height - header_h + (header_h - logo_h) / 2
    try:
        static_folder = app.static_folder or os.path.join(app.root_path, "static")
        logo_rel = app.config.get("LOGO_PRIMARY", "css/lovato_logo.jpg")
        logo_path = os.path.join(static_folder, logo_rel.replace("\\", "/"))
        if os.path.exists(logo_path):
            c.drawImage(logo_path, logo_x, logo_y, width=logo_w, height=logo_h, preserveAspectRatio=True, mask='auto')
            text_x = logo_x + logo_w + 6
        else:
            text_x = x_margin
    except Exception:
        text_x = x_margin

    # Header text
    c.setFillColor(colors.white)
    c.setFont("Helvetica-Bold", 14)
    c.drawString(text_x, height - 11 * mm, str(school_name))
    sub = " ".join(filter(None, [school_address, school_phone, school_email, school_website])) or "Official Payment Receipt"
    c.setFont("Helvetica", 9)
    c.setFillColor(colors.whitesmoke)
    c.drawString(text_x, height - 15.5 * mm, sub)

    # PAID badge on the right
    badge_w, badge_h = 22 * mm, 8 * mm
    badge_x = width - x_margin - badge_w
    badge_y = height - 12 * mm - (badge_h / 2)
    c.setFillColor(colors.white)
    c.setStrokeColor(colors.white)
    c.roundRect(badge_x, badge_y, badge_w, badge_h, 2 * mm, fill=1, stroke=0)
    c.setFillColor(brand_indigo)
    c.setFont("Helvetica-Bold", 8.5)
    c.drawCentredString(badge_x + badge_w / 2, badge_y + 2.6 * mm, "PAID")

    # Content origin (below header)
    y = height - header_h - 10 * mm

    # Amount card
    card_h = 14 * mm
    card_y = y - card_h
    c.setFillColor(light_bg)
    c.setStrokeColor(soft_border)
    c.roundRect(x_margin, card_y, width - 2 * x_margin, card_h, 3 * mm, fill=1, stroke=0)
    c.setFont("Helvetica", 9)
    c.setFillColor(colors.HexColor("#64748b"))
    c.drawCentredString(width / 2, card_y + card_h - 5 * mm, "Amount Paid")
    c.setFillColor(colors.HexColor("#0f172a"))
    c.setFont("Helvetica-Bold", 16)
    c.drawCentredString(width / 2, card_y + 4.8 * mm, f"KES {float(payment.get('amount') or 0):,.2f}")
    y = card_y - content_gap

    # Separator (dashed)
    c.setStrokeColor(colors.lightgrey)
    c.setDash(1, 2)
    c.line(x_margin, y, width - x_margin, y)
    c.setDash()
    y -= content_gap

    # Key-value rows
    def draw_kv(label: str, value: str):
        nonlocal y
        c.setFont("Helvetica", 9)
        c.setFillColor(colors.HexColor("#64748b"))
        c.drawString(x_margin, y, label)
        c.setFillColor(colors.black)
        c.setFont("Helvetica-Bold", 10)
        c.drawRightString(width - x_margin, y, value)
        y -= 6.5 * mm

    pdate = payment.get("date")
    if hasattr(pdate, "strftime"):
        date_str = pdate.strftime("%Y-%m-%d %H:%M")
    else:
        date_str = str(pdate or "N/A")

    draw_kv("Receipt No.", f"#{payment['id']}")
    draw_kv("Date", date_str)
    draw_kv("Student", str(payment.get("student_name") or ""))
    draw_kv("Class", str(payment.get("class_name") or "N/A"))
    draw_kv("Method", str(payment.get("method") or "N/A"))
    draw_kv("Reference", str(payment.get("reference") or "N/A"))
    draw_kv("Year / Term", f"{payment.get('year') or 'N/A'} / {payment.get('term') or 'N/A'}")

    # Secondary separator
    y += 1.5 * mm
    c.setStrokeColor(colors.lightgrey)
    c.setDash(1, 2)
    c.line(x_margin, y, width - x_margin, y)
    c.setDash()
    y -= 7.5 * mm

    draw_kv("Current Balance", f"KES {float(srow['bal'] or 0):,.2f}")
    draw_kv("Credit on Account", f"KES {float(srow['credit'] or 0):,.2f}")

    # Footer
    c.setFillColor(colors.HexColor("#64748b"))
    c.setFont("Helvetica", 9)
    c.drawCentredString(width / 2, max(y, 22 * mm), "Thank you for your payment.")

    # Add QR code for authenticity (signed details)
    try:
        from reportlab.graphics.barcode import qr as rl_qr
        from reportlab.graphics.shapes import Drawing
        from reportlab.graphics import renderPDF
        # Prepare signed JSON payload similar to HTML receipt
        pdate = payment.get("date")
        if hasattr(pdate, "strftime"):
            date_str = pdate.strftime("%Y-%m-%d %H:%M")
        else:
            date_str = str(pdate or "")
        qr_payload = {
            "t": "receipt",
            "rid": int(payment.get("id")),
            "sid": int(payment.get("sid")),
            "amt": round(float(payment.get("amount") or 0.0), 2),
            "cur": "KES",
            "name": payment.get("student_name") or "",
            "cls": payment.get("class_name") or "",
            "dt": date_str,
            "m": payment.get("method") or "",
            "ref": payment.get("reference") or "",
            "term": payment.get("term") or "",
            "year": payment.get("year") or "",
            "school_id": session.get("school_id"),
        }
        canon = json.dumps(qr_payload, sort_keys=True, separators=(",", ":")).encode("utf-8")
        sig = hmac.new(app.secret_key.encode("utf-8"), canon, hashlib.sha256).hexdigest()[:20]
        qr_payload["sig"] = sig
        qr_text = json.dumps(qr_payload, separators=(",", ":"))

        qr_widget = rl_qr.QrCodeWidget(qr_text)
        b = qr_widget.getBounds()
        size = 36 * mm
        w = b[2] - b[0]
        h = b[3] - b[1]
        d = Drawing(size, size)
        qr_widget.transform = [size / w, 0, 0, size / h, 0, 0]
        d.add(qr_widget)
        renderPDF.draw(d, c, width - x_margin - size, 10 * mm)
        c.setFont("Helvetica", 7.5)
        c.setFillColor(colors.HexColor("#64748b"))
        c.drawRightString(width - x_margin, 9 * mm, "Scan for signed receipt data")
    except Exception:
        pass

    c.showPage()
    c.save()
    pdf_bytes = buf.getvalue()
    buf.close()

    return Response(
        pdf_bytes,
        mimetype="application/pdf",
        headers={
            "Content-Disposition": f"attachment; filename=receipt_{payment_id}.pdf",
        },
    )


@app.route("/collections")
def collections_overview():
    from flask import redirect, url_for
    return redirect(url_for('dashboard'))


# ---------- ANALYTICS DATA (LIVE) ----------
@app.route("/api/analytics_data")
def analytics_data():
    """Live analytics payload (charts + KPIs)."""

    db = get_db_connection()
    cursor = db.cursor(dictionary=True)

    # Resolve current academic context
    try:
        cy, ct = get_or_seed_current_term(db)
    except Exception:
        cy, ct = None, None

    # Timeframe filtering (generalized analytics): ignore client timeframe
    # Always compute across all time for aggregate datasets.
    tf = 'all'
    days = None

    # Monthly totals (by first day label for readability)
    if days is not None:
        cursor.execute(
            f"""
            SELECT DATE_FORMAT(MIN(date), '%b %Y') AS month, SUM(amount) AS total
            FROM payments
            WHERE method <> 'Credit Transfer' AND school_id=%s AND date >= (CURRENT_DATE - INTERVAL {days-1} DAY)
            GROUP BY YEAR(date), MONTH(date)
            ORDER BY YEAR(date), MONTH(date)
            """,
            (session.get("school_id"),),
        )
    else:
        if cy and ct in (1, 2, 3):
            cursor.execute(
                """
                SELECT DATE_FORMAT(MIN(date), '%b %Y') AS month, SUM(amount) AS total
                FROM payments
                WHERE method <> 'Credit Transfer' AND school_id=%s AND year=%s AND term=%s
                GROUP BY YEAR(date), MONTH(date)
                ORDER BY YEAR(date), MONTH(date)
                """,
                (session.get("school_id"), cy, ct),
            )
        else:
            cursor.execute(
                """
                SELECT DATE_FORMAT(MIN(date), '%b %Y') AS month, SUM(amount) AS total
                FROM payments
                WHERE method <> 'Credit Transfer' AND school_id=%s
                GROUP BY YEAR(date), MONTH(date)
                ORDER BY YEAR(date), MONTH(date)
                """,
                (session.get("school_id"),),
            )
    monthly_data = cursor.fetchall()

    # Daily trend - timeframe or default last 30 days
    ddays = 364  # show last 12 months of daily trend by default
    cursor.execute(
        f"""
        SELECT DATE(date) AS day, SUM(amount) AS total
        FROM payments
        WHERE date >= (CURRENT_DATE - INTERVAL {ddays} DAY)
          AND method <> 'Credit Transfer' AND school_id=%s
        GROUP BY DATE(date)
        ORDER BY DATE(date)
        """,
        (session.get("school_id"),),
    )
    daily_trend = cursor.fetchall()

    # Class summary
    if days is not None:
        cursor.execute(
            f"""
            SELECT 
                s.class_name,
                COUNT(s.id) AS total_students,
                COALESCE(SUM(p.amount), 0) AS total_paid,
                COALESCE(SUM(COALESCE(s.balance, s.fee_balance)), 0) AS total_pending,
                COALESCE(SUM(s.credit), 0) AS total_credit
            FROM students s
            LEFT JOIN payments p ON s.id = p.student_id AND p.method <> 'Credit Transfer' AND p.school_id=%s AND p.date >= (CURRENT_DATE - INTERVAL {days-1} DAY)
            WHERE s.school_id=%s
            GROUP BY s.class_name
            ORDER BY s.class_name
            """,
            (session.get("school_id"), session.get("school_id")),
        )
    else:
        if cy and ct in (1, 2, 3):
            cursor.execute(
                """
                SELECT 
                    s.class_name,
                    COUNT(s.id) AS total_students,
                    COALESCE(SUM(p.amount), 0) AS total_paid,
                    COALESCE(SUM(COALESCE(s.balance, s.fee_balance)), 0) AS total_pending,
                    COALESCE(SUM(s.credit), 0) AS total_credit
                FROM students s
                LEFT JOIN payments p ON s.id = p.student_id AND p.method <> 'Credit Transfer' AND p.school_id=%s AND p.year=%s AND p.term=%s
                WHERE s.school_id=%s
                GROUP BY s.class_name
                ORDER BY s.class_name
                """,
                (session.get("school_id"), cy, ct, session.get("school_id")),
            )
        else:
            cursor.execute(
                """
                SELECT 
                    s.class_name,
                    COUNT(s.id) AS total_students,
                    COALESCE(SUM(p.amount), 0) AS total_paid,
                    COALESCE(SUM(COALESCE(s.balance, s.fee_balance)), 0) AS total_pending,
                    COALESCE(SUM(s.credit), 0) AS total_credit
                FROM students s
                LEFT JOIN payments p ON s.id = p.student_id AND p.method <> 'Credit Transfer' AND p.school_id=%s
                WHERE s.school_id=%s
                GROUP BY s.class_name
                ORDER BY s.class_name
                """,
                (session.get("school_id"), session.get("school_id")),
            )
    class_summary = cursor.fetchall()

    # Aggregate totals for convenience (avoid client recompute)
    try:
        _rows = class_summary or []
        def _val(r, key, idx):
            try:
                return (r.get(key) if isinstance(r, dict) else r[idx])
            except Exception:
                return 0
        totals_students = int(sum(int(_val(r, "total_students", 1) or 0) for r in _rows))
        totals_collected = float(sum(float(_val(r, "total_paid", 2) or 0) for r in _rows))
        totals_pending = float(sum(float(_val(r, "total_pending", 3) or 0) for r in _rows))
        totals_credit = float(sum(float(_val(r, "total_credit", 4) or 0) for r in _rows))
    except Exception:
        totals_students = 0
        totals_collected = 0.0
        totals_pending = 0.0
        totals_credit = 0.0

    # Payment method breakdown
    if days is not None:
        cursor.execute(
            f"""
            SELECT method, COUNT(*) AS count, COALESCE(SUM(amount), 0) AS total
            FROM payments
            WHERE school_id=%s AND date >= (CURRENT_DATE - INTERVAL {days-1} DAY)
            GROUP BY method
            ORDER BY total DESC
            """,
            (session.get("school_id"),),
        )
    else:
        if cy and ct in (1, 2, 3):
            cursor.execute(
                """
                SELECT method, COUNT(*) AS count, COALESCE(SUM(amount), 0) AS total
                FROM payments
                WHERE school_id=%s AND year=%s AND term=%s
                GROUP BY method
                ORDER BY total DESC
                """,
                (session.get("school_id"), cy, ct),
            )
        else:
            cursor.execute(
                """
                SELECT method, COUNT(*) AS count, COALESCE(SUM(amount), 0) AS total
                FROM payments
                WHERE school_id=%s
                GROUP BY method
                ORDER BY total DESC
                """,
                (session.get("school_id"),),
            )
    method_breakdown = cursor.fetchall()

    # Top debtors (highest balances)
    cursor.execute("SHOW COLUMNS FROM students LIKE 'balance'")
    has_balance = bool(cursor.fetchone())
    balance_col = "balance" if has_balance else "fee_balance"

    cursor.execute(
        f"""
        SELECT name, class_name, COALESCE({balance_col}, 0) AS balance
        FROM students
        WHERE school_id=%s
        ORDER BY COALESCE({balance_col}, 0) DESC
        LIMIT 5
        """,
        (session.get("school_id"),),
    )
    top_debtors = cursor.fetchall()

    # Reminders: count students above a pending threshold (default KES 5,000)
    try:
        import os as _os
        rem_threshold = int((_os.environ.get("REMINDER_MIN_BAL") or "5000").strip())
    except Exception:
        rem_threshold = 5000
    try:
        cursor.execute(
            f"SELECT COUNT(*) AS c FROM students WHERE school_id=%s AND COALESCE({balance_col},0) > %s",
            (session.get("school_id"), rem_threshold),
        )
        reminders_count = int((cursor.fetchone() or {}).get("c", 0))
    except Exception:
        reminders_count = 0

    # Month-over-month change
    if cy and ct in (1, 2, 3):
        cursor.execute(
            """
            SELECT 
                SUM(CASE WHEN YEAR(date) = YEAR(CURRENT_DATE) AND MONTH(date) = MONTH(CURRENT_DATE) THEN amount ELSE 0 END) AS current_month_total,
                SUM(CASE WHEN DATE_FORMAT(date, '%Y-%m') = DATE_FORMAT(DATE_SUB(CURRENT_DATE, INTERVAL 1 MONTH), '%Y-%m') THEN amount ELSE 0 END) AS prev_month_total
            FROM payments
            WHERE method <> 'Credit Transfer' AND school_id=%s AND year=%s AND term=%s
            """,
            (session.get("school_id"), cy, ct),
        )
    else:
        cursor.execute(
            """
            SELECT 
                SUM(CASE WHEN YEAR(date) = YEAR(CURRENT_DATE) AND MONTH(date) = MONTH(CURRENT_DATE) THEN amount ELSE 0 END) AS current_month_total,
                SUM(CASE WHEN DATE_FORMAT(date, '%Y-%m') = DATE_FORMAT(DATE_SUB(CURRENT_DATE, INTERVAL 1 MONTH), '%Y-%m') THEN amount ELSE 0 END) AS prev_month_total
            FROM payments
            WHERE method <> 'Credit Transfer' AND school_id=%s
            """,
            (session.get("school_id"),),
        )
    mom_row = cursor.fetchone() or {"current_month_total": 0, "prev_month_total": 0}
    current_month_total = float(mom_row.get("current_month_total") or 0)
    prev_month_total = float(mom_row.get("prev_month_total") or 0)
    percent_change = (
        round(((current_month_total - prev_month_total) / prev_month_total) * 100, 1)
        if prev_month_total > 0
        else (100.0 if current_month_total > 0 else 0.0)
    )

    # Meta: active classes
    cursor.execute("SELECT COUNT(DISTINCT class_name) AS active_classes FROM students WHERE school_id=%s", (session.get("school_id"),))
    active_classes = (cursor.fetchone() or {}).get("active_classes", 0)

    # Recent payments snapshot (used by analytics widgets)
    recent_payments = []
    try:
        cursor.execute(
            """
            SELECT p.date, s.name, s.class_name, p.amount, p.method
            FROM payments p
            JOIN students s ON s.id = p.student_id
            WHERE p.school_id=%s AND p.method <> 'Credit Transfer'
            ORDER BY p.date DESC
            LIMIT 10
            """,
            (session.get("school_id"),),
        )
        recent_payments = cursor.fetchall() or []
    except Exception:
        recent_payments = []

    db.close()

    # Normalize rows to JSON‑safe primitives (avoid Decimal/date serialization issues)
    try:
        from datetime import date, datetime  # type: ignore
    except Exception:  # pragma: no cover
        date, datetime = None, None

    def _to_float(v):
        try:
            return float(v or 0)
        except Exception:
            return 0.0

    def _to_int(v):
        try:
            return int(v or 0)
        except Exception:
            return 0

    # Monthly totals
    monthly_data = [
        {"month": (r.get("month") if isinstance(r, dict) else r[0]), "total": _to_float((r.get("total") if isinstance(r, dict) else r[1]))}
        for r in (monthly_data or [])
    ]

    # Daily trend with ISO day
    _daily = []
    for r in (daily_trend or []):
        day_val = (r.get("day") if isinstance(r, dict) else r[0])
        if date and isinstance(day_val, (date, datetime)):
            day_str = day_val.isoformat()
        else:
            day_str = str(day_val)
        _daily.append({"day": day_str, "total": _to_float((r.get("total") if isinstance(r, dict) else r[1]))})
    daily_trend = _daily

    # Class summary numbers
    class_summary = [
        {
            "class_name": (r.get("class_name") if isinstance(r, dict) else r[0]),
            "total_students": _to_int((r.get("total_students") if isinstance(r, dict) else r[1])),
            "total_paid": _to_float((r.get("total_paid") if isinstance(r, dict) else r[2])),
            "total_pending": _to_float((r.get("total_pending") if isinstance(r, dict) else r[3])),
            "total_credit": _to_float((r.get("total_credit") if isinstance(r, dict) else r[4])),
        }
        for r in (class_summary or [])
    ]

    # Payment method breakdown
    method_breakdown = [
        {
            "method": (r.get("method") if isinstance(r, dict) else r[0]) or "",
            "count": _to_int((r.get("count") if isinstance(r, dict) else r[1])),
            "total": _to_float((r.get("total") if isinstance(r, dict) else r[2])),
        }
        for r in (method_breakdown or [])
    ]

    # Top debtors
    top_debtors = [
        {
            "name": (r.get("name") if isinstance(r, dict) else r[0]) or "",
            "class_name": (r.get("class_name") if isinstance(r, dict) else r[1]) or "",
            "balance": _to_float((r.get("balance") if isinstance(r, dict) else r[2])),
        }
        for r in (top_debtors or [])
    ]

    # Enrich class summary with percent_paid
    for row in class_summary:
        paid = _to_float(row.get("total_paid"))
        pending = _to_float(row.get("total_pending"))
        total = paid + pending
        row["percent_paid"] = round((paid / total * 100), 1) if total > 0 else 0.0

    # Expected vs Collected for current term (capped by expected to avoid overpay distortion)
    expected_term_total = 0.0
    term_outstanding = 0.0
    paid_capped_sum = 0.0
    term_collection_rate = 0.0
    expected_by_class = {}
    paid_capped_by_class = {}
    try:
        school_id = session.get("school_id")
        if cy and ct in (1, 2, 3) and school_id:
            cur2 = get_db_connection().cursor(dictionary=True)
            try:
                cur2.execute("SELECT id, class_name FROM students WHERE school_id=%s", (school_id,))
                stu = cur2.fetchall() or []
                id2class = {r["id"]: (r.get("class_name") or "") for r in stu}
                ids = list(id2class.keys())
                if ids:
                    def _in_clause(seq):
                        return ",".join(["%s"] * len(seq))
                    # Expected per student for current term (itemized preferred, fallback to legacy term_fees)
                    cur2.execute(
                        f"SELECT student_id, COALESCE(SUM(amount),0) AS tsum FROM student_term_fee_items WHERE year=%s AND term=%s AND student_id IN ({_in_clause(ids)}) GROUP BY student_id",
                        (cy, ct, *ids),
                    )
                    items_map = {r["student_id"]: float(r.get("tsum") or 0) for r in (cur2.fetchall() or [])}
                    cur2.execute(
                        f"SELECT student_id, COALESCE(SUM(fee_amount),0) AS tsum FROM term_fees WHERE year=%s AND term=%s AND student_id IN ({_in_clause(ids)}) GROUP BY student_id",
                        (cy, ct, *ids),
                    )
                    legacy_map = {r["student_id"]: float(r.get("tsum") or 0) for r in (cur2.fetchall() or [])}
                    expected_map = dict(items_map)
                    for sid_i in ids:
                        if sid_i not in expected_map:
                            expected_map[sid_i] = float(legacy_map.get(sid_i) or 0)
                    expected_term_total = float(sum(expected_map.values()) or 0)

                    # Paid per student for term
                    cur2.execute(
                        f"SELECT student_id, COALESCE(SUM(amount),0) AS tsum FROM payments WHERE method <> 'Credit Transfer' AND school_id=%s AND year=%s AND term=%s AND student_id IN ({_in_clause(ids)}) GROUP BY student_id",
                        (school_id, cy, ct, *ids),
                    )
                    paid_map = {r["student_id"]: float(r.get("tsum") or 0) for r in (cur2.fetchall() or [])}

                    # Aggregate per-student to overall and per-class (cap paid at expected)
                    for sid_i in ids:
                        cls = id2class.get(sid_i, "")
                        exp = float(expected_map.get(sid_i) or 0)
                        paidv = float(paid_map.get(sid_i) or 0)
                        term_outstanding += max(exp - paidv, 0.0)
                        cap = min(paidv, exp)
                        paid_capped_sum += cap
                        expected_by_class[cls] = expected_by_class.get(cls, 0.0) + exp
                        paid_capped_by_class[cls] = paid_capped_by_class.get(cls, 0.0) + cap
                    if expected_term_total > 0:
                        term_collection_rate = round((paid_capped_sum / expected_term_total) * 100.0, 1)
            finally:
                try:
                    cur2.close()
                except Exception:
                    pass
    except Exception:
        expected_term_total = 0.0
        term_outstanding = 0.0
        paid_capped_sum = 0.0
        term_collection_rate = 0.0
        expected_by_class = {}
        paid_capped_by_class = {}

    # Shape expected-by-class list for UI
    exp_cls_rows = []
    try:
        all_classes = set(list(expected_by_class.keys()) + list(paid_capped_by_class.keys()))
        for cls in sorted(all_classes):
            exp_v = float(expected_by_class.get(cls, 0.0))
            paid_v = float(paid_capped_by_class.get(cls, 0.0))
            exp_cls_rows.append({
                "class_name": cls or "",
                "expected": exp_v,
                "collected_capped": paid_v,
                "gap": max(exp_v - paid_v, 0.0),
            })
    except Exception:
        exp_cls_rows = []

    # Return the standard, lightweight analytics payload (real-time)
    resp = jsonify(
        {
            "monthly_data": [
                {"month": (r.get("month") if isinstance(r, dict) else r[0]), "total": float((r.get("total") if isinstance(r, dict) else r[1]) or 0)}
                for r in (monthly_data or [])
            ],
            "daily_trend": (lambda rows: (
                [
                    {"day": (r.get("day").isoformat() if hasattr(r.get("day"), "isoformat") else str(r.get("day"))), "total": float((r.get("total") if isinstance(r, dict) else r[1]) or 0)}
                    for r in (rows or [])
                ]
            ))(daily_trend),
            "class_summary": [
                {
                    "class_name": (r.get("class_name") if isinstance(r, dict) else r[0]),
                    "total_students": int((r.get("total_students") if isinstance(r, dict) else r[1]) or 0),
                    "total_paid": float((r.get("total_paid") if isinstance(r, dict) else r[2]) or 0),
                    "total_pending": float((r.get("total_pending") if isinstance(r, dict) else r[3]) or 0),
                    "total_credit": float((r.get("total_credit") if isinstance(r, dict) else r[4]) or 0),
                }
                for r in (class_summary or [])
            ],
            "method_breakdown": [
                {
                    "method": (r.get("method") if isinstance(r, dict) else r[0]) or "",
                    "count": int((r.get("count") if isinstance(r, dict) else r[1]) or 0),
                    "total": float((r.get("total") if isinstance(r, dict) else r[2]) or 0),
                }
                for r in (method_breakdown or [])
            ],
            "top_debtors": [
                {
                    "name": (r.get("name") if isinstance(r, dict) else r[0]) or "",
                    "class_name": (r.get("class_name") if isinstance(r, dict) else r[1]) or "",
                    "balance": float((r.get("balance") if isinstance(r, dict) else r[2]) or 0),
                }
                for r in (top_debtors or [])
            ],
            "recent_payments": [
                {
                    "student": (r.get("name") if isinstance(r, dict) else r[1]) or "Student",
                    "class_name": (r.get("class_name") if isinstance(r, dict) else r[2]) or "Class",
                    "method": (r.get("method") if isinstance(r, dict) else r[4]) or "Method",
                    "amount": float((r.get("amount") if isinstance(r, dict) else r[3]) or 0),
                    "date": (
                        r.get("date").isoformat()
                        if isinstance(r.get("date"), (date, datetime)) and hasattr(r.get("date"), "isoformat")
                        else str(r.get("date") or "")
                    ),
                }
                for r in (recent_payments or [])
            ],
            "mom": {
                "current_month_total": float(current_month_total or 0),
                "prev_month_total": float(prev_month_total or 0),
                "percent_change": float(percent_change or 0),
            },
            "meta": {"active_classes": int(active_classes or 0)},
            "totals": {
                "students": int(totals_students or 0),
                "collected": float(totals_collected or 0),
                "pending": float(totals_pending or 0),
                "credit": float(totals_credit or 0),
            },
        }
    )
    resp.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
    resp.headers["Pragma"] = "no-cache"
    resp.headers["Expires"] = "0"
    db.close()
    return resp

    # Due-date aging using invoices (requires invoices feature)
    due_aging = {
        "not_due": {"count": 0, "amount": 0.0},
        "over_0_30": {"count": 0, "amount": 0.0},
        "over_31_60": {"count": 0, "amount": 0.0},
        "over_61_90": {"count": 0, "amount": 0.0},
        "over_90": {"count": 0, "amount": 0.0},
    }
    payers_timeliness = {
        "paid_on_time": 0.0,
        "paid_late": 0.0,
        "on_time_rate": 0.0,
        "invoices_on_time": 0,
        "invoices_late": 0,
        "invoices_overdue": 0,
    }
    try:
        sid2 = session.get("school_id")
        if sid2 and cy and ct in (1, 2, 3):
            curx = get_db_connection().cursor(dictionary=True)
            try:
                # Invoices for current term and school
                curx.execute(
                    """
                    SELECT i.student_id, i.due_date, i.total
                    FROM invoices i JOIN students s ON s.id = i.student_id
                    WHERE i.year=%s AND i.term=%s AND s.school_id=%s
                    """,
                    (cy, ct, sid2),
                )
                inv_rows = curx.fetchall() or []
                # Paid totals per student for term
                curx.execute(
                    """
                    SELECT student_id, COALESCE(SUM(amount),0) AS total
                    FROM payments
                    WHERE method <> 'Credit Transfer' AND school_id=%s AND year=%s AND term=%s
                    GROUP BY student_id
                    """,
                    (sid2, cy, ct),
                )
                paid_map_term = {r["student_id"]: float(r.get("total") or 0) for r in (curx.fetchall() or [])}

                # Paid on-time vs late per student (join with invoices for due_date)
                curx.execute(
                    """
                    SELECT p.student_id,
                           COALESCE(SUM(CASE WHEN p.date <= i.due_date THEN p.amount ELSE 0 END),0) AS paid_on_time,
                           COALESCE(SUM(CASE WHEN p.date > i.due_date THEN p.amount ELSE 0 END),0) AS paid_late
                    FROM payments p JOIN invoices i
                      ON i.student_id = p.student_id AND i.year = p.year AND i.term = p.term
                    WHERE p.method <> 'Credit Transfer' AND p.school_id=%s AND p.year=%s AND p.term=%s
                    GROUP BY p.student_id
                    """,
                    (sid2, cy, ct),
                )
                tl_rows = {r["student_id"]: {"paid_on_time": float(r.get("paid_on_time") or 0), "paid_late": float(r.get("paid_late") or 0)} for r in (curx.fetchall() or [])}

                # Aggregate buckets and timeliness
                from datetime import date as _date
                today = _date.today()
                total_expected_term = 0.0
                total_on_time = 0.0
                total_late = 0.0
                inv_on_time_count = 0
                inv_late_count = 0
                inv_overdue_count = 0
                for inv in inv_rows:
                    sid_i = inv.get("student_id")
                    total = float(inv.get("total") or 0)
                    total_expected_term += total
                    paid_total = float(paid_map_term.get(sid_i, 0.0))
                    paid_cap = min(paid_total, total)
                    # Due aging bucket based on outstanding
                    out_amt = max(total - paid_cap, 0.0)
                    due = inv.get("due_date")
                    if out_amt > 0 and due is not None:
                        try:
                            d = due if hasattr(due, "toordinal") else due
                            days = (today - d).days
                        except Exception:
                            days = 0
                        if days <= 0:
                            key = "not_due"
                        elif days <= 30:
                            key = "over_0_30"
                        elif days <= 60:
                            key = "over_31_60"
                        elif days <= 90:
                            key = "over_61_90"
                        else:
                            key = "over_90"
                        due_aging[key]["count"] += 1
                        due_aging[key]["amount"] += out_amt
                        if days > 0:
                            inv_overdue_count += 1

                    # Timeliness
                    tl = tl_rows.get(sid_i, {"paid_on_time": 0.0, "paid_late": 0.0})
                    on_time_cap = min(float(tl.get("paid_on_time") or 0), total)
                    total_cap = paid_cap
                    late_cap = max(total_cap - on_time_cap, 0.0)
                    total_on_time += on_time_cap
                    total_late += late_cap
                    # invoice fully paid on time if total_cap>=total and late_cap==0
                    if total_cap >= total and late_cap == 0:
                        inv_on_time_count += 1
                    elif total_cap >= total and late_cap > 0:
                        inv_late_count += 1

                payers_timeliness["paid_on_time"] = total_on_time
                payers_timeliness["paid_late"] = total_late
                payers_timeliness["invoices_on_time"] = inv_on_time_count
                payers_timeliness["invoices_late"] = inv_late_count
                payers_timeliness["invoices_overdue"] = inv_overdue_count
                if total_expected_term > 0:
                    payers_timeliness["on_time_rate"] = round((total_on_time / total_expected_term) * 100.0, 1)
            finally:
                try:
                    curx.close()
                except Exception:
                    pass
    except Exception:
        pass

    # Additional modern analytics
    # 1) Aging by days since last payment (student-level buckets)
    aging = {"0_30": {"count": 0, "amount": 0.0}, "31_60": {"count": 0, "amount": 0.0}, "61_90": {"count": 0, "amount": 0.0}, ">90": {"count": 0, "amount": 0.0}}
    weekday_trend = []  # 1..7 => Sun..Sat
    forecast = {"history": [], "forecast_next": 0.0}
    method_trend = []
    try:
        sid = session.get("school_id")
        if sid:
            cur3 = get_db_connection().cursor(dictionary=True)
            try:
                # Aging by last payment date
                bal_col = "balance"
                try:
                    cur3.execute("SHOW COLUMNS FROM students LIKE 'balance'")
                    if not cur3.fetchone():
                        bal_col = "fee_balance"
                except Exception:
                    bal_col = "fee_balance"
                cur3.execute(
                    f"""
                    SELECT s.id, COALESCE({bal_col},0) AS bal,
                           (SELECT MAX(p.date) FROM payments p WHERE p.student_id=s.id AND p.school_id=%s) AS last_pay
                    FROM students s
                    WHERE s.school_id=%s
                    """,
                    (sid, sid),
                )
                from datetime import date as _date
                today = _date.today()
                for r in (cur3.fetchall() or []):
                    bal = float(r.get("bal") or 0)
                    if bal <= 0:
                        continue
                    lp = r.get("last_pay")
                    days = 9999
                    try:
                        if lp is not None:
                            d = lp.date() if hasattr(lp, "date") else lp
                            days = max(0, (today - d).days)
                    except Exception:
                        pass
                    key = "0_30" if days <= 30 else ("31_60" if days <= 60 else ("61_90" if days <= 90 else ">90"))
                    aging[key]["count"] += 1
                    aging[key]["amount"] += bal

                # Weekday trend of collections (last 26 weeks)
                cur3.execute(
                    """
                    SELECT DAYOFWEEK(date) AS dow, COALESCE(SUM(amount),0) AS total
                    FROM payments
                    WHERE method <> 'Credit Transfer' AND school_id=%s AND date >= (CURRENT_DATE - INTERVAL 182 DAY)
                    GROUP BY DAYOFWEEK(date)
                    ORDER BY dow
                    """,
                    (sid,),
                )
                wmap = {int(r.get("dow") or 0): float(r.get("total") or 0) for r in (cur3.fetchall() or [])}
                weekday_trend = [
                    {"dow": i, "total": float(wmap.get(i, 0.0))} for i in range(1, 8)
                ]

                # Monthly history and simple forecast (next month = avg of last 3 months)
                cur3.execute(
                    """
                    SELECT DATE_FORMAT(date, '%Y-%m') AS ym, COALESCE(SUM(amount),0) AS total
                    FROM payments
                    WHERE method <> 'Credit Transfer' AND school_id=%s AND date >= DATE_SUB(CURRENT_DATE, INTERVAL 12 MONTH)
                    GROUP BY DATE_FORMAT(date, '%Y-%m')
                    ORDER BY ym
                    """,
                    (sid,),
                )
                hist_rows = cur3.fetchall() or []
                forecast["history"] = [{"month": r.get("ym"), "total": float(r.get("total") or 0)} for r in hist_rows]
                last_vals = [float(r.get("total") or 0) for r in hist_rows][-3:]
                if last_vals:
                    forecast["forecast_next"] = float(sum(last_vals) / len(last_vals))

                # Method trend by month (last 6 months)
                cur3.execute(
                    """
                    SELECT DATE_FORMAT(date, '%Y-%m') AS ym, method, COALESCE(SUM(amount),0) AS total
                    FROM payments
                    WHERE method <> 'Credit Transfer' AND school_id=%s AND date >= DATE_SUB(CURRENT_DATE, INTERVAL 6 MONTH)
                    GROUP BY DATE_FORMAT(date, '%Y-%m'), method
                    ORDER BY ym, method
                    """,
                    (sid,),
                )
                method_trend = [
                    {"month": r.get("ym"), "method": (r.get("method") or ""), "total": float(r.get("total") or 0)}
                    for r in (cur3.fetchall() or [])
                ]
            finally:
                try:
                    cur3.close()
                except Exception:
                    pass
    except Exception:
        aging = aging
        weekday_trend = []
        forecast = {"history": [], "forecast_next": 0.0}
        method_trend = []

    total_due = totals_collected + totals_pending
    totals_map = {
        "students": totals_students,
        "collected": totals_collected,
        "pending": totals_pending,
        "credit": totals_credit,
        "collection_rate": (total_due and (totals_collected / total_due) * 100) or 0.0,
        "avg_per_student": (totals_students and (totals_collected / totals_students)) or 0.0,
    }
    resp = jsonify(
        {
            "monthly_data": monthly_data,
            "daily_trend": daily_trend,
            "class_summary": class_summary,
            "method_breakdown": method_breakdown,
            "top_debtors": top_debtors,
            "mom": {
                "current_month_total": current_month_total,
                "prev_month_total": prev_month_total,
                "percent_change": percent_change,
            },
            "meta": {"active_classes": int(active_classes or 0)},
            "reminders": {"threshold": int(rem_threshold), "count": int(reminders_count)},
            "totals": totals_map,
        }
    )
    # Ensure real-time: prevent intermediary/browser caching
    resp.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
    resp.headers["Pragma"] = "no-cache"
    resp.headers["Expires"] = "0"
    return resp


# ---------- DOCUMENTATION ----------
@app.route("/docs")
def docs():
    """Media hub for documentation with featured video support."""
    # Ensure media folder exists
    media_root = os.path.join(app.root_path, "static", "media")
    try:
        os.makedirs(media_root, exist_ok=True)
    except Exception:
        pass

    # Resolve featured video from settings
    featured_name = (get_setting("FEATURED_VIDEO_NAME") or "").strip()
    featured_url = None
    if featured_name:
        candidate = os.path.join(media_root, featured_name)
        if os.path.exists(candidate):
            featured_url = url_for("static", filename=f"media/{featured_name}")

    return render_template("docs.html", featured_url=featured_url, featured_name=featured_name)


@app.route("/docs/media")
def docs_media():
    """List media files under static/media for the library grid."""
    media_root = os.path.join(app.root_path, "static", "media")
    try:
        os.makedirs(media_root, exist_ok=True)
    except Exception:
        pass
    items = []
    allowed = {".mp4", ".webm", ".mov", ".png", ".jpg", ".jpeg", ".gif"}
    try:
        for name in sorted(os.listdir(media_root)):
            ext = os.path.splitext(name)[1].lower()
            if ext not in allowed:
                continue
            mtype = "video" if ext in {".mp4", ".webm", ".mov"} else "image"
            items.append({
                "name": name,
                "type": mtype,
                "url": url_for("static", filename=f"media/{name}"),
            })
    except Exception:
        items = []
    return jsonify({"ok": True, "media": items})


@app.route("/docs/upload", methods=["POST"])
def docs_upload():
    """Upload one or more media files to static/media."""
    media_root = os.path.join(app.root_path, "static", "media")
    try:
        os.makedirs(media_root, exist_ok=True)
    except Exception:
        pass
    files = request.files.getlist("files") if request.files else []
    if not files:
        return jsonify({"ok": False, "error": "No files"}), 400
    allowed = {".mp4", ".webm", ".mov", ".png", ".jpg", ".jpeg", ".gif"}
    saved = []
    for f in files:
        try:
            name = secure_filename(f.filename or "")
            if not name:
                continue
            ext = os.path.splitext(name)[1].lower()
            if ext not in allowed:
                continue
            path = os.path.join(media_root, name)
            f.save(path)
            saved.append(name)
        except Exception:
            continue
    return jsonify({"ok": True, "saved": saved})


@app.route("/docs/media/<name>", methods=["DELETE"])
def docs_media_delete(name: str):
    """Delete a media file by name from static/media."""
    media_root = os.path.join(app.root_path, "static", "media")
    safe_name = secure_filename(name or "")
    if not safe_name:
        return jsonify({"ok": False, "error": "Invalid name"}), 400
    path = os.path.join(media_root, safe_name)
    try:
        if os.path.isfile(path):
            os.remove(path)
        return jsonify({"ok": True})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/docs/feature", methods=["POST"])
def docs_feature():
    """Set the featured video by file name in settings."""
    try:
        data = request.get_json(silent=True) or {}
        name = (data.get("name") or "").strip()
    except Exception:
        name = ""
    name = secure_filename(name)
    if not name:
        return jsonify({"ok": False, "error": "Invalid name"}), 400
    media_root = os.path.join(app.root_path, "static", "media")
    if not os.path.exists(os.path.join(media_root, name)):
        return jsonify({"ok": False, "error": "File not found"}), 404
    try:
        set_school_setting("FEATURED_VIDEO_NAME", name)
        return jsonify({"ok": True})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500



# ---------- REPORTS PAGE ----------
@app.route("/reports")
def reports():
    # Simple page; exports are scoped to current term
    return render_template("reports.html")


# ---------- XLSX EXPORT (optional) ----------
@app.route("/export_fees_xlsx")
def export_fees_xlsx():
    try:
        from openpyxl.workbook import Workbook
    except Exception:
        try:
            flash('XLSX export requires openpyxl. Falling back to ZIP of CSVs.', 'warning')
        except Exception:
            pass
        return redirect(url_for('export_fees_full'))

    db = get_db_connection()
    cur = db.cursor(dictionary=True)
    sid = session.get('school_id')

    # Optional scope params
    year = request.args.get('year')
    term = request.args.get('term')
    start_date = request.args.get('start_date')
    end_date = request.args.get('end_date')
    try:
        y_val = int(year) if (year or '').isdigit() else None
    except Exception:
        y_val = None
    try:
        t_val = int(term) if (term or '').isdigit() else None
    except Exception:
        t_val = None

    # Optional scoping
    year = request.args.get('year')
    term = request.args.get('term')
    start_date = request.args.get('start_date')
    end_date = request.args.get('end_date')
    try:
        y_val = int(year) if (year or '').isdigit() else None
    except Exception:
        y_val = None
    try:
        t_val = int(term) if (term or '').isdigit() else None
    except Exception:
        t_val = None

    # Students sheet: full roster is still useful regardless of scope
    cur.execute("SELECT name, admission_no, class_name, COALESCE(balance, fee_balance) AS balance, COALESCE(credit,0) AS credit FROM students WHERE school_id=%s ORDER BY class_name, name", (sid,))
    students = cur.fetchall() or []

    # Payments sheet (filtered)
    pay_sql = ("SELECT s.name, s.admission_no, s.class_name, p.year, p.term, p.amount, p.method, p.reference, p.date "
               "FROM payments p JOIN students s ON s.id=p.student_id WHERE p.school_id=%s AND p.method <> 'Credit Transfer'")
    params = [sid]
    if y_val and t_val:
        pay_sql += " AND p.year=%s AND p.term=%s"; params += [y_val, t_val]
    else:
        if y_val:
            pay_sql += " AND p.year=%s"; params.append(y_val)
        if start_date:
            pay_sql += " AND p.date >= %s"; params.append(start_date)
        if end_date:
            pay_sql += " AND p.date <= %s"; params.append(end_date)
    pay_sql += " ORDER BY p.date DESC"
    cur.execute(pay_sql, tuple(params))
    payments = cur.fetchall() or []

    # Class summary (roster-based)
    cur.execute("SELECT class_name AS class, COUNT(*) AS total_students, COALESCE(SUM(COALESCE(balance,fee_balance)),0) AS total_pending, COALESCE(SUM(credit),0) AS total_credit FROM students WHERE school_id=%s GROUP BY class_name ORDER BY class_name", (sid,))
    class_summary = cur.fetchall() or []

    # Term summary (respect year filter if provided)
    term_sql = ("SELECT p.year AS year, p.term AS term, COALESCE(SUM(p.amount),0) AS total FROM payments p WHERE p.school_id=%s AND p.method <> 'Credit Transfer'")
    tparams = [sid]
    if y_val:
        term_sql += " AND p.year=%s"; tparams.append(y_val)
    term_sql += " GROUP BY p.year, p.term ORDER BY p.year, p.term"
    cur.execute(term_sql, tuple(tparams))
    term_summary = cur.fetchall() or []

    # Method breakdown (filtered to scope)
    meth_sql = ("SELECT p.method AS method, COUNT(*) AS cnt, COALESCE(SUM(p.amount),0) AS total FROM payments p WHERE p.school_id=%s AND p.method <> 'Credit Transfer'")
    mparams = [sid]
    if y_val and t_val:
        meth_sql += " AND p.year=%s AND p.term=%s"; mparams += [y_val, t_val]
    else:
        if y_val:
            meth_sql += " AND p.year=%s"; mparams.append(y_val)
        if start_date:
            meth_sql += " AND p.date >= %s"; mparams.append(start_date)
        if end_date:
            meth_sql += " AND p.date <= %s"; mparams.append(end_date)
    meth_sql += " GROUP BY p.method ORDER BY total DESC"
    cur.execute(meth_sql, tuple(mparams))
    method_breakdown = cur.fetchall() or []

    db.close()

    from io import BytesIO
    from openpyxl.workbook import Workbook
    wb = Workbook()
    ws = wb.active; ws.title = 'Students'
    ws.append(['Name','Admission No','Class','Balance (KES)','Credit (KES)'])
    for r in students:
        ws.append([r.get('name'), r.get('admission_no'), r.get('class_name'), r.get('balance'), r.get('credit')])

    ws2 = wb.create_sheet('Payments')
    ws2.append(['Student Name','Admission No','Class','Year','Term','Amount (KES)','Method','Reference','Date'])
    for p in payments:
        ws2.append([p.get('name'), p.get('admission_no'), p.get('class_name'), p.get('year'), p.get('term'), p.get('amount'), p.get('method'), p.get('reference'), p.get('date')])

    ws3 = wb.create_sheet('Class Summary')
    ws3.append(['Class','Total Students','Total Pending (KES)','Total Credit (KES)'])
    for c in class_summary:
        ws3.append([c.get('class'), c.get('total_students'), c.get('total_pending'), c.get('total_credit')])

    ws4 = wb.create_sheet('Term Summary')
    ws4.append(['Year','Term','Total Collected (KES)'])
    for t in term_summary:
        ws4.append([t.get('year'), t.get('term'), t.get('total')])

    ws5 = wb.create_sheet('Method Breakdown')
    ws5.append(['Method','Count','Total (KES)'])
    for m in method_breakdown:
        ws5.append([m.get('method'), m.get('cnt'), m.get('total')])

    mem = BytesIO(); wb.save(mem); mem.seek(0)
    from datetime import datetime as _dt
    ts = _dt.now().strftime('%Y%m%d_%H%M%S')
    return Response(mem.getvalue(), mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', headers={'Content-Disposition': f'attachment; filename=fees_full_report_{ts}.xlsx'})

# ---------- RUN ----------

@app.route("/analytics")
@app.route("/my_analytics")
def analytics_dashboard():
    """Analytics dashboard for all boards."""

    db = get_db_connection()
    cursor = db.cursor(dictionary=True)

    # Determine balance column
    cursor.execute("SHOW COLUMNS FROM students LIKE 'balance'")
    has_balance = bool(cursor.fetchone())
    bal_col = "balance" if has_balance else "fee_balance"

    # Totals for KPIs (all-time)
    cursor.execute("SELECT COUNT(*) AS total FROM students WHERE school_id=%s", (session.get("school_id"),))
    total_students = (cursor.fetchone() or {}).get("total", 0)
    cursor.execute(
        "SELECT COALESCE(SUM(amount), 0) AS t FROM payments WHERE method <> 'Credit Transfer' AND school_id=%s",
        (session.get("school_id"),),
    )
    total_collected = float((cursor.fetchone() or {}).get("t", 0) or 0)
    cursor.execute(f"SELECT COALESCE(SUM({bal_col}), 0) AS b FROM students WHERE school_id=%s", (session.get("school_id"),))
    total_balance = float((cursor.fetchone() or {}).get("b", 0) or 0)
    cursor.execute("SELECT COALESCE(SUM(credit), 0) AS c FROM students WHERE school_id=%s", (session.get("school_id"),))
    total_credit = float((cursor.fetchone() or {}).get("c", 0) or 0)

    # Monthly totals (all time)
    cursor.execute(
        """
        SELECT DATE_FORMAT(MIN(date), '%b %Y') AS month, SUM(amount) AS total
        FROM payments
        WHERE method <> 'Credit Transfer' AND school_id=%s
        GROUP BY YEAR(date), MONTH(date)
        ORDER BY YEAR(date), MONTH(date)
        """,
        (session.get("school_id"),),
    )
    monthly_data = cursor.fetchall() or []

    # Daily trend (last 12 months)
    cursor.execute(
        """
        SELECT DATE(date) AS day, SUM(amount) AS total
        FROM payments
        WHERE date >= (CURRENT_DATE - INTERVAL 364 DAY)
          AND method <> 'Credit Transfer' AND school_id=%s
        GROUP BY DATE(date)
        ORDER BY DATE(date)
        """,
        (session.get("school_id"),),
    )
    daily_trend = cursor.fetchall() or []

    # Class summary (all time)
    cursor.execute(
        """
        SELECT 
            s.class_name,
            COUNT(s.id) AS total_students,
            COALESCE(SUM(p.amount), 0) AS total_paid,
            COALESCE(SUM(COALESCE(s.balance, s.fee_balance)), 0) AS total_pending,
            COALESCE(SUM(s.credit), 0) AS total_credit
        FROM students s
        LEFT JOIN payments p ON s.id = p.student_id AND p.method <> 'Credit Transfer' AND p.school_id=%s
        WHERE s.school_id=%s
        GROUP BY s.class_name
        ORDER BY s.class_name
        """,
        (session.get("school_id"), session.get("school_id")),
    )
    class_summary = cursor.fetchall() or []

    # Payment method breakdown (all time)
    cursor.execute(
        """
        SELECT method, COUNT(*) AS count, COALESCE(SUM(amount), 0) AS total
        FROM payments
        WHERE school_id=%s
        GROUP BY method
        ORDER BY total DESC
        """,
        (session.get("school_id"),),
    )
    method_breakdown = cursor.fetchall() or []

    # Top debtors
    cursor.execute(
        f"""
        SELECT name, class_name, COALESCE({bal_col}, 0) AS balance
        FROM students
        WHERE school_id=%s
        ORDER BY COALESCE({bal_col}, 0) DESC
        LIMIT 5
        """,
        (session.get("school_id"),),
    )
    top_debtors = cursor.fetchall() or []

    # Reminders count
    try:
        rem_threshold = int((os.environ.get("REMINDER_MIN_BAL") or "5000").strip())
    except Exception:
        rem_threshold = 5000
    try:
        cursor.execute(
            f"SELECT COUNT(*) AS c FROM students WHERE school_id=%s AND COALESCE({bal_col},0) > %s",
            (session.get("school_id"), rem_threshold),
        )
        reminders_count = int((cursor.fetchone() or {}).get("c", 0))
    except Exception:
        reminders_count = 0

    # Active classes
    try:
        cursor.execute("SELECT COUNT(DISTINCT class_name) AS c FROM students WHERE school_id=%s AND class_name IS NOT NULL", (session.get("school_id"),))
        active_classes = int((cursor.fetchone() or {}).get("c", 0) or 0)
    except Exception:
        active_classes = 0

    # Recent payments (limit to 10 latest entries)
    recent_payments = []
    try:
        cursor.execute(
            """
            SELECT p.date, s.name, s.class_name, p.amount, p.method
            FROM payments p
            JOIN students s ON s.id = p.student_id
            WHERE p.school_id=%s AND p.method <> 'Credit Transfer'
            ORDER BY p.date DESC
            LIMIT 10
            """,
            (session.get("school_id"),),
        )
        recent_payments = cursor.fetchall() or []
    except Exception:
        recent_payments = []

    # MoM change (current vs previous month)
    try:
        cursor.execute(
            "SELECT COALESCE(SUM(amount),0) AS t FROM payments WHERE method <> 'Credit Transfer' AND school_id=%s AND YEAR(date)=YEAR(CURRENT_DATE) AND MONTH(date)=MONTH(CURRENT_DATE)",
            (session.get("school_id"),),
        )
        current_month_total = float((cursor.fetchone() or {}).get("t", 0) or 0)
        cursor.execute(
            "SELECT COALESCE(SUM(amount),0) AS t FROM payments WHERE method <> 'Credit Transfer' AND school_id=%s AND YEAR(date)=YEAR(DATE_SUB(CURRENT_DATE, INTERVAL 1 MONTH)) AND MONTH(date)=MONTH(DATE_SUB(CURRENT_DATE, INTERVAL 1 MONTH))",
            (session.get("school_id"),),
        )
        prev_month_total = float((cursor.fetchone() or {}).get("t", 0) or 0)
        percent_change = (0.0 if prev_month_total == 0 else round(((current_month_total - prev_month_total) / prev_month_total) * 100.0, 1))
    except Exception:
        current_month_total = 0.0
        prev_month_total = 0.0
        percent_change = 0.0

    # Aggregate totals from class_summary as a fallback
    try:
        totals_students = int(sum(int((r.get("total_students") if isinstance(r, dict) else r[1]) or 0) for r in (class_summary or [])))
        totals_collected = float(sum(float((r.get("total_paid") if isinstance(r, dict) else r[2]) or 0) for r in (class_summary or [])))
        totals_pending = float(sum(float((r.get("total_pending") if isinstance(r, dict) else r[3]) or 0) for r in (class_summary or [])))
        totals_credit = float(sum(float((r.get("total_credit") if isinstance(r, dict) else r[4]) or 0) for r in (class_summary or [])))
    except Exception:
        totals_students = total_students
        totals_collected = total_collected
        totals_pending = total_balance
        totals_credit = total_credit

    avg_collection_per_student = total_collected / total_students if total_students else 0.0
    total_due = total_collected + total_balance
    collection_rate_pct = (total_collected / total_due * 100.0) if total_due else 0.0

    def _class_paid_value(record):
        if isinstance(record, dict):
            return float(record.get("total_paid") or 0)
        if isinstance(record, (list, tuple)) and len(record) > 2:
            try:
                return float(record[2] or 0)
            except Exception:
                return 0.0
        return 0.0

    sorted_classes = sorted((class_summary or []), key=_class_paid_value, reverse=True)
    top_classes = sorted_classes[:4]

    analytics_payload = {
        "monthly_data": [
            {"month": (r.get("month") if isinstance(r, dict) else r[0]), "total": float((r.get("total") if isinstance(r, dict) else r[1]) or 0)}
            for r in (monthly_data or [])
        ],
        "daily_trend": [
            {"day": (r.get("day").isoformat() if hasattr(r.get("day"), "isoformat") else str(r.get("day"))), "total": float((r.get("total") if isinstance(r, dict) else r[1]) or 0)}
            for r in (daily_trend or [])
        ],
        "class_summary": [
            {
                "class_name": (r.get("class_name") if isinstance(r, dict) else r[0]),
                "total_students": int((r.get("total_students") if isinstance(r, dict) else r[1]) or 0),
                "total_paid": float((r.get("total_paid") if isinstance(r, dict) else r[2]) or 0),
                "total_pending": float((r.get("total_pending") if isinstance(r, dict) else r[3]) or 0),
                "total_credit": float((r.get("total_credit") if isinstance(r, dict) else r[4]) or 0),
            }
            for r in (class_summary or [])
        ],
        "method_breakdown": [
            {
                "method": (r.get("method") if isinstance(r, dict) else r[0]) or "",
                "count": int((r.get("count") if isinstance(r, dict) else r[1]) or 0),
                "total": float((r.get("total") if isinstance(r, dict) else r[2]) or 0),
            }
            for r in (method_breakdown or [])
        ],
        "top_debtors": [
            {
                "name": (r.get("name") if isinstance(r, dict) else r[0]) or "",
                "class_name": (r.get("class_name") if isinstance(r, dict) else r[1]) or "",
                "balance": float((r.get("balance") if isinstance(r, dict) else r[2]) or 0),
            }
            for r in (top_debtors or [])
        ],
        "recent_payments": [
            {
                "student": (r.get("name") if isinstance(r, dict) else r[1]) or "Student",
                "class_name": (r.get("class_name") if isinstance(r, dict) else r[2]) or "Class",
                "method": (r.get("method") if isinstance(r, dict) else r[4]) or "Method",
                "amount": float((r.get("amount") if isinstance(r, dict) else r[3]) or 0),
                "date": (
                    r.get("date").isoformat()
                    if isinstance(r.get("date"), (datetime,)) and hasattr(r.get("date"), "isoformat")
                    else str(r.get("date") or "")
                ),
            }
            for r in (recent_payments or [])
        ],
        "mom": {
            "current_month_total": float(current_month_total or 0),
            "prev_month_total": float(prev_month_total or 0),
            "percent_change": float(percent_change or 0),
        },
        "meta": {"active_classes": int(active_classes or 0)},
        "reminders": {"threshold": int(rem_threshold), "count": int(reminders_count)},
        "totals": {
            "students": int(totals_students or 0),
            "collected": float(totals_collected or 0),
            "pending": float(totals_pending or 0),
            "credit": float(totals_credit or 0),
        },
    }

    db.close()
    return render_template(
        "analytics.html",
        total_students=total_students,
        total_collected=total_collected,
        total_balance=total_balance,
        total_credit=total_credit,
        analytics_payload=analytics_payload,
        avg_collection_per_student=avg_collection_per_student,
        collection_rate_pct=collection_rate_pct,
        top_classes=top_classes,
    )

from utils.reports_bootstrap import enable_reports_scheduler
try:
    _sched = enable_reports_scheduler(app)
except Exception as _e:
    print("[scheduler] not started:", _e)

if __name__ == "__main__":
    app.run(debug=True)






@app.route("/export_fees_full")
def export_fees_full():
    """Export a comprehensive school fee report as a ZIP with multiple CSVs.

    Includes:
      - students.csv: roster with balances and credit
      - payments.csv: all payments (student, class, year/term, method, ref, date)
      - class_summary.csv: totals per class (students, paid, pending, credit, collection_rate)
      - term_summary.csv: totals per year/term collected
      - method_breakdown.csv: totals per payment method
    """
    import zipfile
    from datetime import datetime
    from io import BytesIO, StringIO

    db = get_db_connection()
    cur = db.cursor(dictionary=True)
    sid = session.get('school_id')

    # Students
    cur.execute(
        """
        SELECT 
            name AS 'Name', 
            admission_no AS 'Admission No', 
            class_name AS 'Class', 
            COALESCE(balance, fee_balance) AS 'Balance (KES)', 
            COALESCE(credit, 0) AS 'Credit (KES)'
        FROM students
        WHERE school_id=%s
        ORDER BY class_name, name
        """,
        (sid,),
    )
    students = cur.fetchall() or []

    # Payments (respect scope)
    pay_sql = (
        "SELECT s.name AS 'Student Name', s.admission_no AS 'Admission No', s.class_name AS 'Class', "
        "p.year AS 'Year', p.term AS 'Term', p.amount AS 'Amount (KES)', p.method AS 'Method', p.reference AS 'Reference', p.date AS 'Date' "
        "FROM payments p JOIN students s ON s.id = p.student_id WHERE p.school_id=%s AND p.method <> 'Credit Transfer'"
    )
    pparams = [sid]
    if y_val and t_val:
        pay_sql += " AND p.year=%s AND p.term=%s"; pparams += [y_val, t_val]
    else:
        if y_val:
            pay_sql += " AND p.year=%s"; pparams.append(y_val)
        if start_date:
            pay_sql += " AND p.date >= %s"; pparams.append(start_date)
        if end_date:
            pay_sql += " AND p.date <= %s"; pparams.append(end_date)
    pay_sql += " ORDER BY p.date DESC"
    cur.execute(pay_sql, tuple(pparams))
    payments = cur.fetchall() or []

    # Class outstanding summary
    cur.execute(
        """
        SELECT class_name AS class,
               COUNT(*) AS total_students,
               COALESCE(SUM(COALESCE(balance, fee_balance)),0) AS total_pending,
               COALESCE(SUM(credit),0) AS total_credit
        FROM students
        WHERE school_id=%s
        GROUP BY class_name
        ORDER BY class_name
        """,
        (sid,),
    )
    class_out = { r['class'] or '': {
        'class': r['class'] or '',
        'total_students': int(r.get('total_students') or 0),
        'total_pending': float(r.get('total_pending') or 0),
        'total_credit': float(r.get('total_credit') or 0),
        'total_paid': 0.0,
    } for r in (cur.fetchall() or []) }

    # Class paid totals from payments (current term if available)
    # Class paid totals from payments. Prefer selected scope; otherwise use current term.
    cy = ct = None
    if y_val and t_val:
        cy, ct = y_val, t_val
    else:
        try:
            cy, ct = get_or_seed_current_term(db)
        except Exception:
            cy, ct = None, None

    if y_val and not t_val:
        cur.execute(
            """
            SELECT s.class_name AS class, COALESCE(SUM(p.amount),0) AS total_paid
            FROM payments p
            JOIN students s ON s.id = p.student_id
            WHERE p.school_id=%s AND p.method <> 'Credit Transfer' AND p.year=%s
            GROUP BY s.class_name
            """,
            (sid, y_val),
        )
    elif cy and ct in (1, 2, 3):
        cur.execute(
            """
            SELECT s.class_name AS class, COALESCE(SUM(p.amount),0) AS total_paid
            FROM payments p
            JOIN students s ON s.id = p.student_id
            WHERE p.school_id=%s AND p.method <> 'Credit Transfer' AND p.year=%s AND p.term=%s
            GROUP BY s.class_name
            """,
            (sid, cy, ct),
        )
    else:
        cur.execute(
            """
            SELECT s.class_name AS class, COALESCE(SUM(p.amount),0) AS total_paid
            FROM payments p
            JOIN students s ON s.id = p.student_id
            WHERE p.school_id=%s AND p.method <> 'Credit Transfer'
            GROUP BY s.class_name
            """,
            (sid,),
        )
    for row in (cur.fetchall() or []):
        c = row.get('class') or ''
        if c not in class_out:
            class_out[c] = {'class': c, 'total_students': 0, 'total_pending': 0.0, 'total_credit': 0.0, 'total_paid': 0.0}
        class_out[c]['total_paid'] = float(row.get('total_paid') or 0)

    # Build expected and capped-paid per class for current term to compute accurate collection rate
    expected_by_class = {}
    paid_capped_by_class = {}
    try:
        # map student->class
        cur = db.cursor(dictionary=True)
        cur.execute("SELECT id, class_name FROM students WHERE school_id=%s", (sid,))
        stu = cur.fetchall() or []
        id2class = {r['id']: (r.get('class_name') or '') for r in stu}
        ids = list(id2class.keys())
        def _in_clause(seq):
            return ",".join(["%s"] * len(seq))
        if ids and cy and ct in (1,2,3):
            # expected per student for term
            cur.execute(
                f"SELECT student_id, COALESCE(SUM(amount),0) AS tsum FROM student_term_fee_items WHERE year=%s AND term=%s AND student_id IN ({_in_clause(ids)}) GROUP BY student_id",
                (cy, ct, *ids),
            )
            items_map = {r['student_id']: float(r.get('tsum') or 0) for r in (cur.fetchall() or [])}
            cur.execute(
                f"SELECT student_id, COALESCE(SUM(fee_amount),0) AS tsum FROM term_fees WHERE year=%s AND term=%s AND student_id IN ({_in_clause(ids)}) GROUP BY student_id",
                (cy, ct, *ids),
            )
            legacy_map = {r['student_id']: float(r.get('tsum') or 0) for r in (cur.fetchall() or [])}
            expected_map = dict(items_map)
            for sid_i in ids:
                if sid_i not in expected_map:
                    expected_map[sid_i] = float(legacy_map.get(sid_i) or 0)
            # paid per student this term
            cur.execute(
                f"SELECT student_id, COALESCE(SUM(amount),0) AS tsum FROM payments WHERE method <> 'Credit Transfer' AND school_id=%s AND year=%s AND term=%s AND student_id IN ({_in_clause(ids)}) GROUP BY student_id",
                (sid, cy, ct, *ids),
            )
            paid_map = {r['student_id']: float(r.get('tsum') or 0) for r in (cur.fetchall() or [])}
            # aggregate per class
            for sid_i in ids:
                cls = id2class.get(sid_i, '')
                exp = float(expected_map.get(sid_i) or 0)
                paid = float(paid_map.get(sid_i) or 0)
                expected_by_class[cls] = expected_by_class.get(cls, 0.0) + exp
                paid_capped_by_class[cls] = paid_capped_by_class.get(cls, 0.0) + min(paid, exp)
    except Exception:
        expected_by_class = {}
        paid_capped_by_class = {}

    class_summary = []
    for k, v in sorted(class_out.items(), key=lambda kv: (kv[0] or '')):
        exp_total = float(expected_by_class.get(k, 0.0))
        paid_capped = float(paid_capped_by_class.get(k, 0.0))
        rate = (paid_capped/exp_total*100.0) if exp_total > 0 else 0.0
        v2 = {
            'Class': v['class'],
            'Total Students': v['total_students'],
            'Total Paid (KES)': round(float(v['total_paid']), 2),
            'Total Pending (KES)': round(float(v['total_pending']), 2),
            'Total Credit (KES)': round(float(v['total_credit']), 2),
            'Collection Rate (%)': round(rate, 1),
        }
        class_summary.append(v2)

    # Term summary (restrict to selected year if provided)
    t_sql = (
        "SELECT p.year AS 'Year', p.term AS 'Term', COALESCE(SUM(p.amount),0) AS 'Total Collected (KES)' "
        "FROM payments p WHERE p.school_id=%s AND p.method <> 'Credit Transfer'"
    )
    t_params = [sid]
    if y_val:
        t_sql += " AND p.year=%s"; t_params.append(y_val)
    t_sql += " GROUP BY p.year, p.term ORDER BY p.year, p.term"
    cur.execute(t_sql, tuple(t_params))
    term_summary = cur.fetchall() or []

    # Method breakdown (respect scope)
    m_sql = (
        "SELECT p.method AS 'Method', COUNT(*) AS 'Count', COALESCE(SUM(p.amount),0) AS 'Total (KES)' "
        "FROM payments p WHERE p.school_id=%s AND p.method <> 'Credit Transfer'"
    )
    m_params = [sid]
    if y_val and t_val:
        m_sql += " AND p.year=%s AND p.term=%s"; m_params += [y_val, t_val]
    else:
        if y_val:
            m_sql += " AND p.year=%s"; m_params.append(y_val)
        if start_date:
            m_sql += " AND p.date >= %s"; m_params.append(start_date)
        if end_date:
            m_sql += " AND p.date <= %s"; m_params.append(end_date)
    m_sql += " GROUP BY p.method ORDER BY 3 DESC"
    cur.execute(m_sql, tuple(m_params))
    method_breakdown = cur.fetchall() or []

    db.close()

    # Build ZIP
    mem = BytesIO()
    with zipfile.ZipFile(mem, 'w', compression=zipfile.ZIP_DEFLATED) as z:
        def _csv_bytes(rows, fieldnames, default_header):
            sio = StringIO()
            fns = fieldnames if rows else default_header
            w = csv.DictWriter(sio, fieldnames=fns)
            w.writeheader()
            if rows:
                w.writerows(rows)
            return sio.getvalue().encode('utf-8')

        z.writestr('students.csv', _csv_bytes(students, (students[0].keys() if students else []), ['Name','Admission No','Class','Balance (KES)','Credit (KES)']))
        z.writestr('payments.csv', _csv_bytes(payments, (payments[0].keys() if payments else []), ['Student Name','Admission No','Class','Year','Term','Amount (KES)','Method','Reference','Date']))
        z.writestr('class_summary.csv', _csv_bytes(class_summary, (class_summary[0].keys() if class_summary else []), ['Class','Total Students','Total Paid (KES)','Total Pending (KES)','Total Credit (KES)','Collection Rate (%)']))
        z.writestr('term_summary.csv', _csv_bytes(term_summary, (term_summary[0].keys() if term_summary else []), ['Year','Term','Total Collected (KES)']))
        z.writestr('method_breakdown.csv', _csv_bytes(method_breakdown, (method_breakdown[0].keys() if method_breakdown else []), ['Method','Count','Total (KES)']))

    mem.seek(0)
    ts = datetime.now().strftime('%Y%m%d_%H%M%S')
    return Response(mem.getvalue(), mimetype='application/zip', headers={'Content-Disposition': f'attachment; filename=fees_full_report_{ts}.zip'})


 



# ---------- SCHOOL PROFILE DOCX ----------
@app.route("/export_school_profile_docx")
def export_school_profile_docx():

    try:
        from docx import Document
        from docx.shared import Pt, Inches
    except Exception:
        try:
            flash('Word export requires python-docx. Please install python-docx.', 'warning')
        except Exception:
            pass
        return redirect(url_for('reports'))

    db = get_db_connection()
    cur = db.cursor(dictionary=True)
    sid = session.get('school_id')

    # Optional filters
    year = request.args.get('year')
    term = request.args.get('term')
    start_date = request.args.get('start_date')
    end_date = request.args.get('end_date')
    try:
        y_val = int(year) if (year or '').isdigit() else None
    except Exception:
        y_val = None
    try:
        t_val = int(term) if (term or '').isdigit() else None
    except Exception:
        t_val = None

    school_name = get_setting('SCHOOL_NAME') or get_setting('APP_NAME') or 'School'
    school_address = get_setting('SCHOOL_ADDRESS') or ''
    school_phone = get_setting('SCHOOL_PHONE') or ''
    school_email = get_setting('SCHOOL_EMAIL') or ''
    school_website = get_setting('SCHOOL_WEBSITE') or ''
    from datetime import datetime as _dt
    generated_at = _dt.now()
    generated_str = generated_at.strftime('%Y-%m-%d %H:%M:%S')

    cur.execute("SELECT COUNT(*) AS total FROM students WHERE school_id=%s", (sid,))
    total_students = (cur.fetchone() or {}).get('total', 0)

    # Totals collected can be constrained by date/year/term if provided
    if y_val and t_val:
        cur.execute("SELECT COALESCE(SUM(amount),0) AS total FROM payments WHERE method <> 'Credit Transfer' AND school_id=%s AND year=%s AND term=%s", (sid, y_val, t_val))
    else:
        cur.execute("SELECT COALESCE(SUM(amount),0) AS total FROM payments WHERE method <> 'Credit Transfer' AND school_id=%s", (sid,))
    total_collected = (cur.fetchone() or {}).get('total', 0)

    cur.execute("SHOW COLUMNS FROM students LIKE 'balance'")
    has_balance = bool(cur.fetchone())
    column = 'balance' if has_balance else 'fee_balance'
    cur.execute(f"SELECT COALESCE(SUM({column}),0) AS total FROM students WHERE school_id=%s", (sid,))
    total_balance = (cur.fetchone() or {}).get('total', 0)
    cur.execute("SELECT COALESCE(SUM(credit),0) AS total FROM students WHERE school_id=%s", (sid,))
    total_credit = (cur.fetchone() or {}).get('total', 0)

    auth_qr_data = ""
    try:
        qr_payload = {
            "t": "school_profile",
            "sid": int(sid or 0),
            "s": school_name,
            "col": round(float(total_collected or 0), 2),
            "bal": round(float(total_balance or 0), 2),
            "students": int(total_students or 0),
            "ts": generated_at.strftime('%Y-%m-%dT%H:%M:%S'),
        }
        canon = json.dumps(qr_payload, sort_keys=True, separators=(",", ":")).encode("utf-8")
        sig = hmac.new(app.secret_key.encode("utf-8"), canon, hashlib.sha256).hexdigest()[:20]
        qr_payload["sig"] = sig
        auth_qr_data = json.dumps(qr_payload, separators=(",", ":"))
    except Exception:
        auth_qr_data = ""

    # Class summary
    cur.execute("""
        SELECT class_name AS class,
               COUNT(*) AS total_students,
               COALESCE(SUM(COALESCE(balance, fee_balance)),0) AS total_pending,
               COALESCE(SUM(credit),0) AS total_credit
        FROM students WHERE school_id=%s GROUP BY class_name ORDER BY class_name
    """, (sid,))
    class_rows = cur.fetchall() or []

    # Paid per class (date/term filters if present)
    paid_sql = "SELECT s.class_name AS class, COALESCE(SUM(p.amount),0) AS total_paid FROM payments p JOIN students s ON s.id=p.student_id WHERE p.school_id=%s AND p.method <> 'Credit Transfer'"
    params = [sid]
    if y_val and t_val:
        paid_sql += " AND p.year=%s AND p.term=%s"; params += [y_val, t_val]
    if start_date:
        paid_sql += " AND p.date >= %s"; params.append(start_date)
    if end_date:
        paid_sql += " AND p.date <= %s"; params.append(end_date)
    paid_sql += " GROUP BY s.class_name"
    cur.execute(paid_sql, tuple(params))
    paid_map = {r['class']: float(r.get('total_paid') or 0) for r in (cur.fetchall() or [])}

    # Top debtors (independent of filters)
    cur.execute("SELECT name, class_name, COALESCE(balance, fee_balance) AS balance FROM students WHERE school_id=%s ORDER BY balance DESC LIMIT 10", (sid,))
    top_debtors = cur.fetchall() or []

    # Recent payments (with optional date filters)
    r_sql = "SELECT p.date, s.name, s.class_name, p.amount, p.method, p.reference FROM payments p JOIN students s ON s.id=p.student_id WHERE p.school_id=%s"
    r_params = [sid]
    if y_val and t_val:
        r_sql += " AND p.year=%s AND p.term=%s"; r_params += [y_val, t_val]
    if start_date:
        r_sql += " AND p.date >= %s"; r_params.append(start_date)
    if end_date:
        r_sql += " AND p.date <= %s"; r_params.append(end_date)
    r_sql += " ORDER BY p.date DESC LIMIT 10"
    cur.execute(r_sql, tuple(r_params))
    recent_payments = cur.fetchall() or []

    # Monthly totals & method breakdown (with filters)
    m_sql = "SELECT DATE_FORMAT(date, '%Y-%m') AS ym, SUM(amount) AS total FROM payments WHERE method <> 'Credit Transfer' AND school_id=%s"
    m_params = [sid]
    if y_val and t_val:
        m_sql += " AND year=%s AND term=%s"; m_params += [y_val, t_val]
    if start_date:
        m_sql += " AND date >= %s"; m_params.append(start_date)
    if end_date:
        m_sql += " AND date <= %s"; m_params.append(end_date)
    m_sql += " GROUP BY DATE_FORMAT(date, '%Y-%m') ORDER BY ym DESC LIMIT 12"
    cur.execute(m_sql, tuple(m_params))
    monthly_rows = list(reversed(cur.fetchall() or []))

    md_sql = "SELECT method, SUM(amount) AS total FROM payments WHERE method <> 'Credit Transfer' AND school_id=%s"
    md_params = [sid]
    if y_val and t_val:
        md_sql += " AND year=%s AND term=%s"; md_params += [y_val, t_val]
    if start_date:
        md_sql += " AND date >= %s"; md_params.append(start_date)
    if end_date:
        md_sql += " AND date <= %s"; md_params.append(end_date)
    md_sql += " GROUP BY method ORDER BY total DESC"
    cur.execute(md_sql, tuple(md_params))
    method_rows = cur.fetchall() or []

    # Top credit
    cur.execute("SELECT name, class_name, credit FROM students WHERE school_id=%s AND credit > 0 ORDER BY credit DESC LIMIT 10", (sid,))
    top_credit = cur.fetchall() or []

    # Expected vs Collected per term (filters if provided)
    if y_val and t_val:
        cur.execute("SELECT sti.year AS year, sti.term AS term, COALESCE(SUM(sti.amount),0) AS total FROM student_term_fee_items sti JOIN students s ON s.id = sti.student_id WHERE s.school_id=%s AND sti.year=%s AND sti.term=%s GROUP BY sti.year, sti.term", (sid, y_val, t_val))
    else:
        cur.execute("SELECT sti.year AS year, sti.term AS term, COALESCE(SUM(sti.amount),0) AS total FROM student_term_fee_items sti JOIN students s ON s.id = sti.student_id WHERE s.school_id=%s GROUP BY sti.year, sti.term", (sid,))
    items_rows = cur.fetchall() or []

    if y_val and t_val:
        cur.execute("SELECT tf.year AS year, tf.term AS term, COALESCE(SUM(tf.fee_amount),0) AS total FROM term_fees tf JOIN students s ON s.id = tf.student_id WHERE s.school_id=%s AND tf.year=%s AND tf.term=%s GROUP BY tf.year, tf.term", (sid, y_val, t_val))
    else:
        cur.execute("SELECT tf.year AS year, tf.term AS term, COALESCE(SUM(tf.fee_amount),0) AS total FROM term_fees tf JOIN students s ON s.id = tf.student_id WHERE s.school_id=%s GROUP BY tf.year, tf.term", (sid,))
    legacy_rows = cur.fetchall() or []

    exp_map = {}
    for r in items_rows:
        key = (int(r.get('year') or 0), int(r.get('term') or 0)); exp_map[key] = float(r.get('total') or 0)
    for r in legacy_rows:
        key = (int(r.get('year') or 0), int(r.get('term') or 0)); exp_map[key] = exp_map.get(key,0.0) + float(r.get('total') or 0)

    c_sql = "SELECT year, term, COALESCE(SUM(amount),0) AS total FROM payments WHERE method <> 'Credit Transfer' AND school_id=%s"
    c_params = [sid]
    if y_val and t_val:
        c_sql += " AND year=%s AND term=%s"; c_params += [y_val, t_val]
    if start_date:
        c_sql += " AND date >= %s"; c_params.append(start_date)
    if end_date:
        c_sql += " AND date <= %s"; c_params.append(end_date)
    c_sql += " GROUP BY year, term ORDER BY year, term"
    cur.execute(c_sql, tuple(c_params))
    coll_rows = cur.fetchall() or []

    exp_table = []
    for r in coll_rows:
        key = (int(r.get('year') or 0), int(r.get('term') or 0))
        exp = float(exp_map.get(key, 0.0))
        col = float(r.get('total') or 0)
        exp_table.append({'year': key[0], 'term': key[1], 'expected': exp, 'collected': col, 'gap': exp - col})

    db.close()

    doc = Document()
    # Add a reusable header with school profile (logo + contacts)
    try:
        from docx.shared import Inches
        from docx.enum.table import WD_TABLE_ALIGNMENT
        import os
        section = doc.sections[0]
        header = section.header
        table = header.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        left, right = table.rows[0].cells
        # Logo (left)
        try:
            logo_rel = get_setting('SCHOOL_LOGO_URL') or app.config.get('LOGO_PRIMARY')
            if logo_rel:
                lp = logo_rel
                if not os.path.isabs(lp):
                    lp = os.path.join(app.root_path, 'static', logo_rel)
                if os.path.exists(lp):
                    left.paragraphs[0].add_run().add_picture(lp, width=Inches(1.0))
        except Exception:
            pass
        # School identity (right)
        rp = right.paragraphs[0]
        r1 = rp.add_run(str(school_name))
        r1.bold = True
        r1.font.size = Pt(14)
        rp.add_run("\n")
        sub_parts = [p for p in [school_address, school_phone, school_email, get_setting('SCHOOL_WEBSITE') or ''] if p]
        if sub_parts:
            rp.add_run(" • ".join(sub_parts)).font.size = Pt(10)
    except Exception:
        pass

    # Document title and metadata
    doc.add_heading("Fee Report", 0)
    doc.add_paragraph(f"Generated: {generated_str}")
    doc.add_heading("School Profile", level=1)
    profile_data = [
        ("School Name", school_name),
        ("Address", school_address or "Not provided"),
        ("Phone", school_phone or "Not provided"),
        ("Email", school_email or "Not provided"),
    ]
    if school_website:
        profile_data.append(("Website", school_website))
    profile_table = doc.add_table(rows=len(profile_data), cols=2)
    profile_table.style = "Table Grid"
    for idx, (label, value) in enumerate(profile_data):
        cells = profile_table.rows[idx].cells
        cells[0].text = label
        cells[1].text = value or "Not provided"

    # Determine collection rate for a specific term using expected fees per student (exclude overpay)
    try:
        if y_val and t_val:
            y_sel, t_sel = int(y_val), int(t_val)
        else:
            y_sel, t_sel = get_or_seed_current_term(get_db_connection())
    except Exception:
        y_sel, t_sel = None, None
    expected_term_total = float(exp_map.get((int(y_sel or 0), int(t_sel or 0)), 0.0)) if (y_sel and t_sel) else 0.0
    collected_term_total = 0.0
    for rr in (coll_rows or []):
        if int(rr.get('year') or 0) == int(y_sel or -1) and int(rr.get('term') or 0) == int(t_sel or -1):
            collected_term_total = float(rr.get('total') or 0)
            break
    try:
        rate = round((min(collected_term_total, expected_term_total) / expected_term_total) * 100.0, 1) if expected_term_total > 0 else 0.0
    except Exception:
        rate = 0.0
    doc.add_heading('Summary', level=1)
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light List'
    labels = ['Total Students','Total Collected (KES)','Total Pending (KES)','Total Credit (KES)','Collection Rate (%)']
    values = [total_students, total_collected, total_balance, total_credit, round(rate,1)]
    for i,(k,v) in enumerate(zip(labels, values)):
        table.cell(i,0).text = str(k)
        table.cell(i,1).text = f"{v}"

    doc.add_heading('Class Summary', level=1)
    ct = doc.add_table(rows=1, cols=6); ct.style = 'Light Grid'
    for j,h in enumerate(['Class','Total Students','Paid (KES)','Pending (KES)','Credit (KES)','Rate %']): ct.rows[0].cells[j].text = h
    # Compute expected and capped-paid per class for selected/current term
    expected_by_class = {}
    paid_capped_by_class = {}
    try:
        cur2 = get_db_connection().cursor(dictionary=True)
        cur2.execute("SELECT id, class_name FROM students WHERE school_id=%s", (sid,))
        stu = cur2.fetchall() or []
        id2class = {r['id']: (r.get('class_name') or '') for r in stu}
        ids = list(id2class.keys())
        def _in_clause(seq):
            return ",".join(["%s"] * len(seq))
        if ids and y_sel and t_sel:
            cur2.execute(
                f"SELECT student_id, COALESCE(SUM(amount),0) AS tsum FROM student_term_fee_items WHERE year=%s AND term=%s AND student_id IN ({_in_clause(ids)}) GROUP BY student_id",
                (y_sel, t_sel, *ids),
            )
            items_map = {r['student_id']: float(r.get('tsum') or 0) for r in (cur2.fetchall() or [])}
            cur2.execute(
                f"SELECT student_id, COALESCE(SUM(fee_amount),0) AS tsum FROM term_fees WHERE year=%s AND term=%s AND student_id IN ({_in_clause(ids)}) GROUP BY student_id",
                (y_sel, t_sel, *ids),
            )
            legacy_map = {r['student_id']: float(r.get('tsum') or 0) for r in (cur2.fetchall() or [])}
            expected_map = dict(items_map)
            for sid_i in ids:
                if sid_i not in expected_map:
                    expected_map[sid_i] = float(legacy_map.get(sid_i) or 0)
            cur2.execute(
                f"SELECT student_id, COALESCE(SUM(amount),0) AS tsum FROM payments WHERE method <> 'Credit Transfer' AND school_id=%s AND year=%s AND term=%s AND student_id IN ({_in_clause(ids)}) GROUP BY student_id",
                (sid, y_sel, t_sel, *ids),
            )
            paid_map_by_stu = {r['student_id']: float(r.get('tsum') or 0) for r in (cur2.fetchall() or [])}
            for sid_i in ids:
                cls = id2class.get(sid_i, '')
                exp = float(expected_map.get(sid_i) or 0)
                paidv = float(paid_map_by_stu.get(sid_i) or 0)
                expected_by_class[cls] = expected_by_class.get(cls, 0.0) + exp
                paid_capped_by_class[cls] = paid_capped_by_class.get(cls, 0.0) + min(paidv, exp)
    except Exception:
        expected_by_class = {}
        paid_capped_by_class = {}

    for row in class_rows:
        cls = row.get('class') or ''
        paid = float(paid_map.get(cls, 0))
        pending = float(row.get('total_pending') or 0)
        credit = float(row.get('total_credit') or 0)
        exp = float(expected_by_class.get(cls, 0.0))
        paid_cap = float(paid_capped_by_class.get(cls, 0.0))
        rate = round((paid_cap/exp*100.0),1) if exp>0 else 0.0
        r = ct.add_row().cells
        r[0].text = str(cls)
        r[1].text = str(row.get('total_students') or 0)
        r[2].text = str(round(paid,2))
        r[3].text = str(round(pending,2))
        r[4].text = str(round(credit,2))
        r[5].text = str(rate)

    doc.add_heading('Top Debtors', level=1)
    tdt = doc.add_table(rows=1, cols=3); tdt.style = 'Light Grid'
    for j,h in enumerate(['Student','Class','Balance (KES)']): tdt.rows[0].cells[j].text = h
    for d in (top_debtors or []):
        r = tdt.add_row().cells
        r[0].text = str(d.get('name') or '')
        r[1].text = str(d.get('class_name') or '')
        r[2].text = str(round(float(d.get('balance') or 0),2))

    doc.add_heading('Top Credit', level=1)
    tct = doc.add_table(rows=1, cols=3); tct.style = 'Light Grid'
    for j,h in enumerate(['Student','Class','Credit (KES)']): tct.rows[0].cells[j].text = h
    for c in (top_credit or []):
        r = tct.add_row().cells
        r[0].text = str(c.get('name') or '')
        r[1].text = str(c.get('class_name') or '')
        r[2].text = str(round(float(c.get('credit') or 0),2))

    doc.add_heading('Recent Payments', level=1)
    rpt = doc.add_table(rows=1, cols=6); rpt.style = 'Light Grid'
    for j,h in enumerate(['Date','Student','Class','Amount','Method','Ref']): rpt.rows[0].cells[j].text = h
    for pmt in (recent_payments or []):
        r = rpt.add_row().cells
        r[0].text = str(pmt.get('date') or '')
        r[1].text = str(pmt.get('name') or '')
        r[2].text = str(pmt.get('class_name') or '')
        r[3].text = str(round(float(pmt.get('amount') or 0),2))
        r[4].text = str(pmt.get('method') or '')
        r[5].text = str(pmt.get('reference') or '')

    doc.add_heading('Expected vs Collected (by term)', level=1)
    ect = doc.add_table(rows=1, cols=5); ect.style = 'Light Grid'
    for j,h in enumerate(['Year','Term','Expected (KES)','Collected (KES)','Gap (KES)']): ect.rows[0].cells[j].text = h
    for r in (exp_table or []):
        row = ect.add_row().cells
        row[0].text = str(r.get('year') or '')
        row[1].text = str(r.get('term') or '')
        row[2].text = str(round(float(r.get('expected') or 0),2))
        row[3].text = str(round(float(r.get('collected') or 0),2))
        row[4].text = str(round(float(r.get('gap') or 0),2))

    try:
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt
        from io import BytesIO as _BIO
        if monthly_rows:
            fig, ax = plt.subplots(figsize=(6,2.2))
            labels = [r.get('ym') for r in monthly_rows]
            vals = [float(r.get('total') or 0) for r in monthly_rows]
            ax.plot(labels, vals, marker='o', color='#4f46e5'); ax.set_title('Monthly Collections (last 12)')
            ax.tick_params(axis='x', rotation=45, labelsize=7); ax.grid(alpha=.2)
            buf = _BIO(); plt.tight_layout(); fig.savefig(buf, format='png', dpi=200); plt.close(fig); buf.seek(0)
            doc.add_picture(buf, width=Inches(6))
        if method_rows:
            fig, ax = plt.subplots(figsize=(5,3.2))
            labels = [str(r.get('method') or 'N/A') for r in method_rows]
            vals = [float(r.get('total') or 0) for r in method_rows]
            if sum(vals) <= 0:
                vals = [1]
                labels = ['No Data']
            wedges, texts, autotexts = ax.pie(
                vals,
                labels=None,  # avoid label clutter on the chart
                autopct=lambda p: f"{p:.0f}%" if p >= 5 else "",
                startangle=90,
                pctdistance=0.72,
                labeldistance=1.15,
                wedgeprops={"linewidth": 1, "edgecolor": "#fff"},
            )
            ax.set_title('Method Breakdown')
            # Legend to the right for readable labels
            ax.legend(wedges, labels, title='Method', loc='center left', bbox_to_anchor=(1.02, 0.5), frameon=False)
            buf2 = _BIO(); plt.tight_layout(); fig.savefig(buf2, format='png', dpi=200, bbox_inches='tight'); plt.close(fig); buf2.seek(0)
            doc.add_picture(buf2, width=Inches(5.5))
    except Exception:
        pass

    if auth_qr_data:
        doc.add_heading("Document Authentication", level=1)
        doc.add_paragraph("Scan the QR code to verify this school profile export and its signed metadata.")
        try:
            import qrcode
            from io import BytesIO as _QRBIO

            qr = qrcode.QRCode(
                version=3,
                error_correction=qrcode.constants.ERROR_CORRECT_Q,
                box_size=8,
                border=2,
            )
            qr.add_data(auth_qr_data)
            qr.make(fit=True)
            img = qr.make_image(fill_color="black", back_color="white")
            qr_buf = _QRBIO()
            img.save(qr_buf, format="PNG")
            qr_buf.seek(0)
            doc.add_picture(qr_buf, width=Inches(1.9))
        except Exception:
            pass

    from io import BytesIO
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    ts = _dt.now().strftime('%Y%m%d_%H%M%S')
    return Response(bio.getvalue(), mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', headers={'Content-Disposition': f'attachment; filename=school_profile_{ts}.docx'})
